from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime, timezone
from hashlib import sha256
import json
import os
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
from typing import Any, Iterable, Sequence
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import pandas as pd
from lxml import etree
from pypdf import PdfReader

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
DCTERMS_NS = "http://purl.org/dc/terms/"

A4_WIDTH = Inches(8.27)
A4_HEIGHT = Inches(11.69)
MARGIN_LEFT = Inches(1.06)
MARGIN_RIGHT = Inches(0.95)
MARGIN_TOP = Inches(0.95)
MARGIN_BOTTOM = Inches(0.87)
USABLE_WIDTH_INCHES = 8.27 - 1.06 - 0.95
HEADING_1_COLOR = "365F91"
HEADING_2_COLOR = "4F81BD"
TABLE_HEADER_FILL = "D9EAF7"


def _sha256(path: Path) -> str:
    digest = sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _xml_bytes(root: etree._Element) -> bytes:
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def create_sanitized_template(reference_docx: Path, template_docx: Path) -> None:
    """Create a content-free template while retaining the reference style system."""

    reference_docx = reference_docx.resolve()
    template_docx = template_docx.resolve()
    template_docx.parent.mkdir(parents=True, exist_ok=True)
    expected_hash = "06bdda2fa0fcb377ce8789c9e336db707115f8566dd5bd09ffb7bb7eb11b22b5"
    actual_hash = _sha256(reference_docx)
    if actual_hash.lower() != expected_hash:
        raise ValueError(f"参考 Word 哈希发生变化：{actual_hash}")

    removed_prefixes = (
        "word/media/",
        "customXml/",
        "word/comments",
        "word/people",
        "word/threadedComments",
    )
    removed_exact = {
        "docProps/thumbnail.jpeg",
        "docProps/custom.xml",
    }

    with ZipFile(reference_docx, "r") as source, ZipFile(template_docx, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            name = item.filename
            if name in removed_exact or name.startswith(removed_prefixes):
                continue
            data = source.read(name)

            if name == "word/document.xml":
                root = etree.fromstring(data)
                body = root.find(f"{{{WORD_NS}}}body")
                if body is None:
                    raise ValueError("参考 Word 缺少 document body")
                sect_pr = body.find(f"{{{WORD_NS}}}sectPr")
                for child in list(body):
                    body.remove(child)
                body.append(etree.Element(f"{{{WORD_NS}}}p"))
                if sect_pr is not None:
                    body.append(sect_pr)
                for element in root.iter():
                    for attr in list(element.attrib):
                        if attr.startswith(f"{{{WORD_NS}}}rsid"):
                            del element.attrib[attr]
                data = _xml_bytes(root)

            elif name == "word/_rels/document.xml.rels":
                root = etree.fromstring(data)
                for rel in list(root):
                    rel_type = rel.get("Type", "")
                    if any(
                        rel_type.endswith(suffix)
                        for suffix in ("/image", "/customXml", "/comments", "/commentsExtended")
                    ):
                        root.remove(rel)
                data = _xml_bytes(root)

            elif name == "_rels/.rels":
                root = etree.fromstring(data)
                for rel in list(root):
                    rel_type = rel.get("Type", "")
                    if rel_type.endswith("/metadata/thumbnail") or rel_type.endswith("/custom-properties"):
                        root.remove(rel)
                data = _xml_bytes(root)

            elif name == "[Content_Types].xml":
                root = etree.fromstring(data)
                for override in list(root):
                    part = override.get("PartName", "")
                    if part.startswith("/customXml/") or part.startswith("/word/comments"):
                        root.remove(override)
                    if part == "/docProps/custom.xml":
                        root.remove(override)
                data = _xml_bytes(root)

            elif name == "docProps/core.xml":
                root = etree.fromstring(data)
                for xpath in (
                    f"{{{DC_NS}}}creator",
                    f"{{{CP_NS}}}lastModifiedBy",
                    f"{{{DC_NS}}}title",
                    f"{{{DC_NS}}}subject",
                    f"{{{DC_NS}}}description",
                    f"{{{CP_NS}}}keywords",
                    f"{{{CP_NS}}}category",
                ):
                    element = root.find(xpath)
                    if element is not None:
                        element.text = ""
                title = root.find(f"{{{DC_NS}}}title")
                if title is not None:
                    title.text = "最优化理论与算法课程报告模板"
                for xpath in (f"{{{DCTERMS_NS}}}created", f"{{{DCTERMS_NS}}}modified"):
                    element = root.find(xpath)
                    if element is not None:
                        root.remove(element)
                revision = root.find(f"{{{CP_NS}}}revision")
                if revision is not None:
                    revision.text = "1"
                data = _xml_bytes(root)

            target.writestr(item, data)

    # The template must remain a valid DOCX after package-level sanitization.
    check = Document(template_docx)
    if len(check.sections) != 1:
        raise ValueError("清洁模板的节结构异常")


def _set_style_font(style: Any, east_asia: str, latin: str, size: float, *, bold: bool | None = None,
                    color: str | None = None) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key, value in (("w:ascii", latin), ("w:hAnsi", latin), ("w:eastAsia", east_asia), ("w:cs", latin)):
        r_fonts.set(qn(key), value)


def _set_run_font(run: Any, east_asia: str = "宋体", latin: str = "Times New Roman",
                  size: float = 11.0, *, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = latin
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key, value in (("w:ascii", latin), ("w:hAnsi", latin), ("w:eastAsia", east_asia), ("w:cs", latin)):
        r_fonts.set(qn(key), value)


def _set_widow_control(paragraph: Any, enabled: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    widow = p_pr.find(qn("w:widowControl"))
    if widow is None:
        widow = OxmlElement("w:widowControl")
        p_pr.append(widow)
    widow.set(qn("w:val"), "1" if enabled else "0")


def _configure_styles(doc: DocumentObject) -> None:
    normal = doc.styles["Normal"]
    _set_style_font(normal, "宋体", "Times New Roman", 11.0)
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)

    h1 = doc.styles["Heading 1"]
    _set_style_font(h1, "宋体", "Times New Roman", 15.0, bold=True, color=HEADING_1_COLOR)
    h1.paragraph_format.first_line_indent = Pt(0)
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    _set_style_font(h2, "宋体", "Times New Roman", 12.0, bold=True, color=HEADING_2_COLOR)
    h2.paragraph_format.first_line_indent = Pt(0)
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    _set_style_font(h3, "宋体", "Times New Roman", 11.0, bold=True, color=HEADING_2_COLOR)
    h3.paragraph_format.first_line_indent = Pt(0)
    h3.paragraph_format.space_before = Pt(4)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    _set_style_font(caption, "宋体", "Times New Roman", 10.5)
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.line_spacing = 1.15

    if "Equation" not in [style.name for style in doc.styles]:
        equation = doc.styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = doc.styles["Equation"]
    _set_style_font(equation, "宋体", "Cambria Math", 11.0)
    equation.paragraph_format.first_line_indent = Pt(0)
    equation.paragraph_format.space_before = Pt(4)
    equation.paragraph_format.space_after = Pt(6)
    equation.paragraph_format.keep_together = True

    styles_by_lower_name = {style.name.lower(): style for style in doc.styles}
    for name, left in (("toc 1", 0), ("toc 2", 16)):
        if name in styles_by_lower_name:
            style = styles_by_lower_name[name]
            _set_style_font(style, "宋体", "Times New Roman", 10.5)
            style.paragraph_format.left_indent = Pt(left)
            style.paragraph_format.first_line_indent = Pt(0)
            style.paragraph_format.line_spacing = 1.0
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)


def _configure_section(section: Any) -> None:
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)


def _clear_story(story: Any) -> Any:
    paragraph = story.paragraphs[0]
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    for extra in list(story.paragraphs[1:]):
        extra._element.getparent().remove(extra._element)
    return paragraph


def _set_page_numbering(section: Any, fmt: str, start: int) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    pg_num.set(qn("w:start"), str(start))


def _append_field(paragraph: Any, instruction: str, result: str, *, bookmark: str | None = None,
                  bookmark_id: int | None = None, font_size: float = 10.5) -> None:
    if bookmark is not None and bookmark_id is not None:
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), bookmark)
        paragraph._p.append(start)

    begin_run = paragraph.add_run()
    _set_run_font(begin_run, size=font_size)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    _set_run_font(instruction_run, size=font_size)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    instruction_run._r.append(instr)

    separate_run = paragraph.add_run()
    _set_run_font(separate_run, size=font_size)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = paragraph.add_run(result)
    _set_run_font(result_run, size=font_size)

    end_run = paragraph.add_run()
    _set_run_font(end_run, size=font_size)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)

    if bookmark is not None and bookmark_id is not None:
        finish = OxmlElement("w:bookmarkEnd")
        finish.set(qn("w:id"), str(bookmark_id))
        paragraph._p.append(finish)


def _add_page_number(section: Any, fmt: str, start: int) -> None:
    section.footer.is_linked_to_previous = False
    paragraph = _clear_story(section.footer)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    _append_field(paragraph, "PAGE", "1", font_size=10.0)
    _set_page_numbering(section, fmt, start)


def _set_update_fields(doc: DocumentObject) -> None:
    settings = doc.settings._element
    element = settings.find(qn("w:updateFields"))
    if element is None:
        element = OxmlElement("w:updateFields")
        settings.append(element)
    element.set(qn("w:val"), "true")


def _add_title_paragraph(doc: DocumentObject, text: str, *, size: float, before: float,
                         after: float = 0, bold: bool = True) -> Any:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    _set_run_font(run, "宋体", "Times New Roman", size, bold=bold)
    return paragraph


def _add_body(doc: DocumentObject, text: str, *, first_indent: bool = True,
              keep_with_next: bool = False) -> Any:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(21 if first_indent else 0)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    _set_widow_control(paragraph)
    run = paragraph.add_run(text)
    _set_run_font(run)
    return paragraph


def _add_body_with_ref(doc: DocumentObject, prefix: str, bookmark: str, suffix: str,
                       cached_result: str) -> Any:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(21)
    _set_widow_control(paragraph)
    run = paragraph.add_run(prefix)
    _set_run_font(run)
    _append_field(paragraph, f"REF {bookmark} \\h", cached_result, font_size=11.0)
    run = paragraph.add_run(suffix)
    _set_run_font(run)
    return paragraph


def _add_heading(doc: DocumentObject, level: int, text: str) -> Any:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    _set_widow_control(paragraph)
    size, color = {
        1: (15.0, HEADING_1_COLOR),
        2: (12.0, HEADING_2_COLOR),
        3: (11.0, HEADING_2_COLOR),
    }[level]
    for run in paragraph.runs:
        _set_run_font(
            run,
            "宋体",
            "Times New Roman",
            size,
            bold=True,
            color=color,
        )
    return paragraph


def _add_list_item(doc: DocumentObject, text: str) -> Any:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.first_line_indent = None
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        _set_run_font(run)
    if not paragraph.runs:
        run = paragraph.add_run(text)
        _set_run_font(run)
    else:
        paragraph.runs[0].text = text
    return paragraph


def _add_front_title(doc: DocumentObject, text: str) -> Any:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run(text)
    _set_run_font(run, "宋体", "Times New Roman", 16.0, bold=True)
    return paragraph


def _add_toc(doc: DocumentObject) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    _append_field(paragraph, 'TOC \\o "1-1" \\h \\z \\u', "目录将在 Microsoft Word 中更新", font_size=11.0)


def _find_mml2omml() -> Path:
    candidates = [
        Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"),
        Path(r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    raise FileNotFoundError("未找到 Microsoft Office MML2OMML.XSL")


class EquationTransformer:
    def __init__(self) -> None:
        self._transform = etree.XSLT(etree.parse(str(_find_mml2omml())))

    def convert(self, mathml: str) -> etree._Element:
        root = etree.fromstring(mathml.encode("utf-8"))
        converted = self._transform(root)
        element = deepcopy(converted.getroot())
        if element.tag == f"{{{MATH_NS}}}oMathPara":
            child = element.find(f"{{{MATH_NS}}}oMath")
            if child is None:
                raise ValueError("MathML 转换未生成 oMath")
            return deepcopy(child)
        return element


def _add_equation(doc: DocumentObject, transformer: EquationTransformer, mathml: str, number: int) -> None:
    paragraph = doc.add_paragraph(style="Equation")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    p_pr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    center_tab = OxmlElement("w:tab")
    center_tab.set(qn("w:val"), "center")
    center_tab.set(qn("w:pos"), str(int(USABLE_WIDTH_INCHES * 1440 / 2)))
    right_tab = OxmlElement("w:tab")
    right_tab.set(qn("w:val"), "right")
    right_tab.set(qn("w:pos"), str(int(USABLE_WIDTH_INCHES * 1440)))
    tabs.extend([center_tab, right_tab])
    p_pr.append(tabs)
    paragraph.add_run("\t")
    paragraph._p.append(transformer.convert(mathml))
    run = paragraph.add_run(f"\t（{number}）")
    _set_run_font(run, size=10.5)


def _set_cell_margins(cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths_inches: Sequence[float]) -> None:
    widths = [int(value * 1440) for value in widths_inches]
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _format_table(table: Any, widths_inches: Sequence[float], *, font_size: float = 9.5) -> None:
    _set_table_geometry(table, widths_inches)
    header_row = table.rows[0]
    tr_pr = header_row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)

    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                tc_pr = cell._tc.get_or_add_tcPr()
                shading = tc_pr.find(qn("w:shd"))
                if shading is None:
                    shading = OxmlElement("w:shd")
                    tc_pr.append(shading)
                shading.set(qn("w:fill"), TABLE_HEADER_FILL)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    _set_run_font(run, size=font_size, bold=True if row_index == 0 else None)


@dataclass
class FieldState:
    figure: int = 0
    table: int = 0
    bookmark_id: int = 100


def _add_table_caption(doc: DocumentObject, state: FieldState, title: str) -> str:
    state.table += 1
    state.bookmark_id += 1
    bookmark = f"tbl{state.table}"
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run("表 ")
    _set_run_font(run, size=10.5)
    _append_field(
        paragraph,
        "SEQ Table \\* ARABIC",
        str(state.table),
        bookmark=bookmark,
        bookmark_id=state.bookmark_id,
        font_size=10.5,
    )
    run = paragraph.add_run(f"　{title}")
    _set_run_font(run, size=10.5)
    return bookmark


def _add_table(doc: DocumentObject, state: FieldState, title: str, headers: Sequence[str],
               rows: Iterable[Sequence[str]], widths_inches: Sequence[float], *,
               font_size: float = 9.5) -> str:
    bookmark = _add_table_caption(doc, state, title)
    data_rows = list(rows)
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = str(value)
    for row_index, values in enumerate(data_rows, start=1):
        for column_index, value in enumerate(values):
            table.rows[row_index].cells[column_index].text = str(value)
    _format_table(table, widths_inches, font_size=font_size)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.first_line_indent = Pt(0)
    return bookmark


def _add_figure(doc: DocumentObject, state: FieldState, image_path: Path, title: str,
                *, width: float = 6.0) -> str:
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    state.figure += 1
    state.bookmark_id += 1
    bookmark = f"fig{state.figure}"
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.first_line_indent = Pt(0)
    image_paragraph.paragraph_format.space_before = Pt(4)
    image_paragraph.paragraph_format.space_after = Pt(0)
    image_paragraph.paragraph_format.keep_with_next = True
    run = image_paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    for doc_pr in run._r.xpath(".//wp:docPr"):
        doc_pr.set("descr", title)
        doc_pr.set("title", title)

    caption = doc.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_together = True
    prefix = caption.add_run("图 ")
    _set_run_font(prefix, size=10.5)
    _append_field(
        caption,
        "SEQ Figure \\* ARABIC",
        str(state.figure),
        bookmark=bookmark,
        bookmark_id=state.bookmark_id,
        font_size=10.5,
    )
    suffix = caption.add_run(f"　{title}")
    _set_run_font(suffix, size=10.5)
    return bookmark


def _pct(value: float, digits: int = 2) -> str:
    return f"{float(value):.{digits}%}"


def _num(value: float, digits: int = 2) -> str:
    return f"{float(value):.{digits}f}"


def _sci(value: float, digits: int = 2) -> str:
    return f"{float(value):.{digits}e}"


def _load_frame(tables_dir: Path, filename: str) -> pd.DataFrame:
    return pd.read_csv(tables_dir / filename)


def _strategy_label(name: str) -> str:
    return {
        "equal_weight": "等权组合",
        "inverse_downside_vol": "逆下行波动率",
        "erc": "风险平价",
    }.get(name, name)


def _estimator_label(name: str) -> str:
    return {
        "sample": "样本协方差",
        "ewma_full": "EWMA 全协方差",
        "ewma_semi": "EWMA 半协方差",
    }.get(name, name)


def _method_label(name: str) -> str:
    return {
        "newton": "阻尼牛顿法",
        "lbfgsb": "L-BFGS-B",
        "slsqp": "SLSQP",
    }.get(name, name)


def _set_document_properties(doc: DocumentObject, config: dict[str, Any]) -> None:
    props = doc.core_properties
    props.title = str(config["title"])
    props.subject = "最优化理论与算法课程报告"
    props.author = ""
    props.last_modified_by = ""
    props.keywords = "风险平价；凸优化；阻尼牛顿法；风险预算；资产配置"
    props.comments = "由项目数据、数值实验和样本外回测结果可复现生成"
    report_date = str(config.get("report_date", ""))
    match = re.search(r"(\d{4})\D+(\d{1,2})(?:\D+(\d{1,2}))?", report_date)
    if match:
        year, month, day = (int(match.group(1)), int(match.group(2)), int(match.group(3) or 1))
        own_timestamp = datetime(year, month, day, tzinfo=timezone.utc)
        props.created = own_timestamp
        props.modified = own_timestamp
    props.revision = 1


def _clear_document_body(doc: DocumentObject) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _add_cover(doc: DocumentObject, config: dict[str, Any]) -> None:
    _add_title_paragraph(doc, "最优化理论与算法课程报告", size=22, before=110)
    cover_title = str(config["title"]).replace("——", "\n——", 1)
    _add_title_paragraph(doc, cover_title, size=18, before=45, after=10)
    _add_title_paragraph(doc, f"课程名称：{config['course']}", size=14, before=18, bold=False)
    _add_title_paragraph(doc, f"姓名：{config['student_name']}", size=14, before=18, bold=False)
    _add_title_paragraph(doc, f"学号：{config['student_id']}", size=14, before=18, bold=False)
    _add_title_paragraph(doc, str(config.get("report_date", "2026年7月")), size=12, before=55, bold=False)


def _add_abstract(doc: DocumentObject, config: dict[str, Any], summary: dict[str, Any]) -> None:
    validation = summary["validation_erc"]
    newton = summary["newton_summary"]
    _add_front_title(doc, "摘　要")
    abstract = (
        "风险平价通过配置风险而非直接配置资本，在缺少稳定收益预测时为多资产组合提供一种可解释的分散化方法。"
        "本文以等风险贡献（ERC）为研究对象，首先依据 Euler 分解建立资产风险贡献，并将风险预算方程重构为正权重域上的"
        "对数障碍凸优化问题；随后推导解析梯度、Hessian 与 KKT 条件，说明最优解的唯一性以及凸模型与目标风险贡献之间的"
        "等价关系。在算法层面，本文实现带正权重可行步长和 Armijo 回溯的阻尼牛顿法，并以最大绝对风险贡献偏差不超过"
        " 10⁻⁶ 作为独立验收标准。148 个滚动 EWMA 半协方差矩阵上的数值实验表明，阻尼牛顿法中位迭代次数为 "
        f"{newton['median_iterations']:.0f}，中位风险贡献误差为 {_sci(newton['median_rc_error'])}，在精度和迭代效率上"
        "优于本文采用的通用算法基准。基于 9 类 ETF 或指数代理资产的样本外检验显示，风险平价组合年化收益为 "
        f"{_pct(validation['annual_return'])}、年化波动率为 {_pct(validation['annual_volatility'])}、夏普比率为 "
        f"{_num(validation['sharpe'])}、最大回撤为 {_pct(validation['max_drawdown'])}。结果表明，其主要价值在于降低波动和"
        "回撤，而非保证最高绝对收益。最后，非等额风险预算实验验证了同一凸模型对一般风险预算的推广能力。"
    )
    _add_body(doc, abstract)
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(12)
    lead = paragraph.add_run("关键词：")
    _set_run_font(lead, "宋体", "Times New Roman", 11.0, bold=True)
    rest = paragraph.add_run("风险平价；风险预算；凸优化；阻尼牛顿法；资产配置")
    _set_run_font(rest)


def _add_report_content_legacy(doc: DocumentObject, course_dir: Path, config: dict[str, Any],
                               summary: dict[str, Any]) -> None:
    tables_dir = course_dir / "output" / "tables"
    figures_dir = course_dir / "output" / "figures"
    data_quality = summary["data_quality"]
    validation_erc = summary["validation_erc"]
    validation_equal = summary["validation_equal_weight"]
    optimizer = _load_frame(tables_dir, "optimizer_evolution_summary.csv").sort_values("stage_order")
    algorithms = _load_frame(tables_dir, "algorithm_summary.csv")
    stress = _load_frame(tables_dir, "stress_test_summary.csv")
    strategies = _load_frame(tables_dir, "strategy_metrics.csv")
    estimators = _load_frame(tables_dir, "estimator_comparison.csv")
    representative = _load_frame(tables_dir, "representative_optimal_solution.csv")
    risk_budget = _load_frame(tables_dir, "risk_budget_extension.csv")
    sensitivity = _load_frame(tables_dir, "parameter_sensitivity.csv")
    selected = summary["selected_parameter"]
    rb_summary = summary["risk_budget_extension"]
    newton = summary["newton_summary"]
    state = FieldState()
    equations = EquationTransformer()

    # 一、引言
    _add_heading(doc, 1, "一、引言")
    _add_heading(doc, 2, "（一）研究背景")
    _add_body(
        doc,
        "现代投资组合理论通常以 Markowitz 均值—方差框架为起点，通过在预期收益与方差之间进行权衡来确定资产权重[1]。"
        "然而，预期收益在有限样本中估计误差较大，协方差矩阵也会随市场状态发生变化。对于包含股票、债券、商品和黄金的"
        "多资产组合，如果仅按资金权重平均配置，高波动资产往往会主导组合总风险，使名义上的分散投资并未形成真正的风险分散。"
        "风险平价由此将优化目标从资本比例转向风险贡献比例，在不依赖精确收益预测的前提下，把每个资产对总风险的贡献控制在"
        "给定预算附近[2-4]。"
    )
    _add_body(
        doc,
        "风险贡献相等并不意味着资金权重相等。低波动、低相关资产通常需要更高资金权重，才能与高波动资产贡献相近的组合风险。"
        "这一机制使风险平价具有直观的分散化含义，但也把问题转化为带正权重约束的非线性方程组。直接最小化风险贡献差异时，"
        "目标函数量级会随收益单位和协方差尺度显著变化，通用求解器可能因默认停止条件而在近似等权的初始点提前终止。"
        "因此，本研究重点并非简单比较多个策略的历史收益，而是考察模型重构、尺度处理、二阶信息和解后验收如何共同提高求解可靠性。"
    )
    _add_heading(doc, 2, "（二）相关研究与课程问题")
    _add_body(
        doc,
        "Maillard、Roncalli 与 Teïletche 系统讨论了等风险贡献组合的性质，并指出 ERC 可视为风险预算模型的重要特例[2]。"
        "Bruder 与 Roncalli 将风险平价推广到任意正风险预算，强调预算约束描述的是风险暴露而非资金仓位[3]；Roncalli 的专著"
        "进一步给出了风险预算的建模体系和资产配置应用[4]。在算法层面，Spinu 证明风险预算问题可通过严格凸目标求解，并提出"
        "基于 Newton 方法的高效算法[5]。这些研究与凸优化、数值优化课程中的严格凸性、KKT 条件、Newton 方向和线搜索理论"
        "直接对应[6-7]。"
    )
    _add_body(
        doc,
        "本项目的历史代码恰好记录了从经验性修补到理论重构的过程：早期版本使用原始风险贡献平方差和 SLSQP；随后通过目标放大"
        "和相对误差处理数值尺度；再通过对数障碍目标获得严格凸模型；当前版本进一步使用解析 Hessian、阻尼牛顿方向、"
        "正权重步长与 Armijo 回溯。课程报告据此回答三个问题：第一，哪些历史变化属于优化器改进，哪些只是数据或风险估计变化；"
        "第二，凸重构和二阶方法为何能提升求解精度；第三，这些数值改进是否能在统一实验口径下被验证。"
    )
    _add_heading(doc, 2, "（三）研究贡献与结构")
    _add_body(
        doc,
        "本文的第一项贡献是建立可审计的历史优化器复现框架，把 v0.02 至 v0.05 与当前阻尼牛顿法放在相同的 148 个风险矩阵上比较，"
        "并把 SciPy 的 success 标志与真正的风险贡献验收分开。第二项贡献是从目标函数、梯度、Hessian 与 KKT 条件出发，说明"
        "ERC 与一般风险预算的凸等价关系。第三项贡献是把动态可行资产集合、一般风险预算和外部波动率覆盖区分为不同建模层级，"
        "从而避免把所有策略变化笼统称为优化器升级。"
    )

    # 二、数据与实验设计
    _add_heading(doc, 1, "二、数据、风险估计与实验设计")
    _add_heading(doc, 2, "（一）数据来源与资产范围")
    _add_body(
        doc,
        f"实验使用项目原始 Excel 工作簿的第 {data_quality['sheet_index'] + 1} 个工作表，样本从 "
        f"{data_quality['start_date']} 延伸至 {data_quality['end_date']}，共 {data_quality['rows']} 个交易日、"
        f"{data_quality['assets']} 类资产。资产池覆盖沪深 300、中证 1000、红利低波、10 年和 30 年国债、有色、能源化工、"
        "豆粕与黄金等 ETF 或指数代理。ETF 成立前使用项目已提供的指数代理，因此结果刻画的是统一资产类别暴露，而不是所有时期"
        "均可真实交易的同一只基金。"
    )
    data_rows = [
        ("样本区间", f"{data_quality['start_date']}—{data_quality['end_date']}"),
        ("交易日数量", str(data_quality["rows"])),
        ("资产数量", str(data_quality["assets"])),
        ("重复日期", str(data_quality["duplicate_dates"])),
        ("原始缺失单元格", str(data_quality["missing_cells_before_fill"])),
        ("收益率范围", f"{_pct(data_quality['min_return'])} 至 {_pct(data_quality['max_return'])}"),
    ]
    data_table = _add_table(
        doc, state, "样本与数据质量概览", ["项目", "统计结果"], data_rows, [2.0, 4.18], font_size=10.0
    )
    _add_body_with_ref(
        doc,
        "数据质量检查结果见表 ",
        data_table,
        "。原始数据只有 3 个孤立缺失收益，按项目既定口径填 0；收益率原表为百分数，进入模型前统一除以 100。"
        "v0.02 对收益单位的修复会改变协方差量级和历史绩效，但并未改变求解器目标、约束或停止条件，因此不能被解释为优化器升级。",
        str(state.table),
    )
    _add_heading(doc, 2, "（二）EWMA 半协方差与正则化")
    _add_body(
        doc,
        "RiskMetrics 技术文档推动了指数加权移动平均方法在风险估计中的应用[8]。本文采用 252 个交易日滚动窗口和衰减系数 0.97，"
        "对较近观测赋予更高权重；同时仅保留收益的负向部分构造下行二阶矩，以便风险度量更关注亏损状态。由于半协方差、短样本"
        "和高度相关资产可能导致矩阵接近奇异，v0.06 及当前实现对角加入 10⁻⁸I。该操作改变的是风险矩阵的数值条件，而不是"
        "风险预算目标本身。"
    )
    ewma_math = f"""<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mi>Σ</mi><mo>=</mo><munderover><mo>∑</mo><mrow><mi>t</mi><mo>=</mo><mn>1</mn></mrow><mi>T</mi></munderover><msub><mi>α</mi><mi>t</mi></msub><msub><mi>r</mi><mi>t</mi></msub><msubsup><mi>r</mi><mi>t</mi><mi>T</mi></msubsup><mo>+</mo><msup><mn>10</mn><mrow><mo>−</mo><mn>8</mn></mrow></msup><mi>I</mi><mo>,</mo><mspace width="0.5em"/><msub><mi>α</mi><mi>t</mi></msub><mo>∝</mo><msup><mn>{config['decay']}</mn><mrow><mi>T</mi><mo>−</mo><mi>t</mi></mrow></msup></mrow></math>"""
    _add_equation(doc, equations, ewma_math, 1)
    _add_heading(doc, 2, "（三）滚动回测与受控比较")
    _add_body(
        doc,
        f"组合在每个月末使用截至当日的 {config['window']} 日历史窗口估计风险矩阵，目标权重在下一交易日执行，单边换手成本为 "
        f"{config['fee_rate'] * 10000:.0f}bp。训练期为 {config['train_start']} 至 {config['train_end']}，验证期为 "
        f"{config['validation_start']} 至 {config['validation_end']}。所有参数选择只使用训练期，样本外区间不参与调参；"
        "该时间顺序由无前视测试验证。"
    )
    _add_body(
        doc,
        "优化器演进实验固定风险矩阵、风险预算、初始点和验收标准，仅改变历史目标函数及相应求解设置。每个版本输出求解器状态、"
        "迭代次数、运行时间、最大绝对风险贡献误差、权重和误差及最小权重。最终正确性统一由风险贡献误差和约束误差判断，"
        "不把通用求解器返回 success=True 直接等同于结果正确。运行时间受硬件、Python 与底层线性代数库影响，只作为同一环境内参考。"
    )

    # 三、模型
    _add_heading(doc, 1, "三、ERC 模型与凸等价重构")
    _add_heading(doc, 2, "（一）组合风险与风险贡献")
    _add_body(
        doc,
        "设 n 个资产的资金权重为 w，协方差矩阵为 Σ，且 wᵢ≥0、∑wᵢ=1。组合方差由二次型给出。"
        "对组合波动率求偏导可得到边际风险贡献；将边际贡献与资金权重相乘并按组合总风险归一化，得到各资产风险贡献比例。"
    )
    _add_equation(
        doc,
        equations,
        """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msubsup><mi>σ</mi><mi>p</mi><mn>2</mn></msubsup><mo stretchy="false">（</mo><mi>w</mi><mo stretchy="false">）</mo><mo>=</mo><msup><mi>w</mi><mi>T</mi></msup><mi>Σ</mi><mi>w</mi></mrow></math>""",
        2,
    )
    _add_equation(
        doc,
        equations,
        """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msub><mi>RC</mi><mi>i</mi></msub><mo>=</mo><mfrac><mrow><msub><mi>w</mi><mi>i</mi></msub><msub><mrow><mo stretchy="false">（</mo><mi>Σ</mi><mi>w</mi><mo stretchy="false">）</mo></mrow><mi>i</mi></msub></mrow><mrow><msup><mi>w</mi><mi>T</mi></msup><mi>Σ</mi><mi>w</mi></mrow></mfrac><mo>,</mo><mspace width="0.5em"/><munderover><mo>∑</mo><mrow><mi>i</mi><mo>=</mo><mn>1</mn></mrow><mi>n</mi></munderover><msub><mi>RC</mi><mi>i</mi></msub><mo>=</mo><mn>1</mn></mrow></math>""",
        3,
    )
    _add_body(
        doc,
        "等风险贡献组合要求 RCᵢ=1/n；一般风险预算则要求 RCᵢ=bᵢ，其中 bᵢ>0 且 ∑bᵢ=1。"
        "早期代码直接最小化实际贡献与目标贡献之差的平方和。该目标易于理解，但相对于协方差尺度并非无量纲，"
        "在日收益或百分数收益之间切换时，函数值和梯度量级会发生数量级变化。"
    )
    _add_equation(
        doc,
        equations,
        """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><munderover><mo>min</mo><mrow><mi>w</mi><mo>≥</mo><mn>0</mn></mrow><mrow><mn>1</mn><mo>ᵀ</mo><mi>w</mi><mo>=</mo><mn>1</mn></mrow></munderover><mspace width="0.4em"/><munderover><mo>∑</mo><mrow><mi>i</mi><mo>=</mo><mn>1</mn></mrow><mi>n</mi></munderover><msup><mrow><mo stretchy="false">（</mo><msub><mi>RC</mi><mi>i</mi></msub><mo>−</mo><msub><mi>b</mi><mi>i</mi></msub><mo stretchy="false">）</mo></mrow><mn>2</mn></msup></mrow></math>""",
        4,
    )
    _add_heading(doc, 2, "（二）对数障碍凸目标")
    _add_body(
        doc,
        "v0.05 不再直接优化风险贡献差异，而是在正正交象限上求解二次项与对数障碍项之和。对任意正定 Σ 和正风险预算 b，"
        "二次项是严格凸函数，−log xᵢ 同样严格凸，因此总目标存在唯一极小点[5-6]。由于目标对整体尺度并不要求权重和为 1，"
        "求得正向量 x 后再归一化为 w=x/(1ᵀx)，即可得到满足预算风险贡献的资金权重。"
    )
    _add_equation(
        doc,
        equations,
        """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><munderover><mo>min</mo><mi>x</mi><mrow><mi>x</mi><mo>&gt;</mo><mn>0</mn></mrow></munderover><mspace width="0.4em"/><mi>f</mi><mo stretchy="false">（</mo><mi>x</mi><mo stretchy="false">）</mo><mo>=</mo><mfrac><mn>1</mn><mn>2</mn></mfrac><msup><mi>x</mi><mi>T</mi></msup><mi>Σ</mi><mi>x</mi><mo>−</mo><munderover><mo>∑</mo><mrow><mi>i</mi><mo>=</mo><mn>1</mn></mrow><mi>n</mi></munderover><msub><mi>b</mi><mi>i</mi></msub><mi>log</mi><msub><mi>x</mi><mi>i</mi></msub></mrow></math>""",
        5,
    )
    _add_heading(doc, 2, "（三）梯度、Hessian 与 KKT 条件")
    _add_body(
        doc,
        "解析一阶与二阶导数既用于证明严格凸性，也为 Newton 方法提供计算基础。Hessian 由正定矩阵 Σ 与正对角矩阵叠加而成，"
        "因此在可行域内始终正定，Newton 方向唯一。最优点的一阶条件 Σx−b⊘x=0；逐元素乘以 x 后得到"
        " xᵢ(Σx)ᵢ=bᵢ，归一化并利用 Euler 分解即可得到 RCᵢ=bᵢ。由此可见，凸模型并非近似 ERC，"
        "而是风险预算方程在正权重域中的等价求解形式。"
    )
    _add_equation(
        doc,
        equations,
        """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mo>∇</mo><mi>f</mi><mo stretchy="false">（</mo><mi>x</mi><mo stretchy="false">）</mo><mo>=</mo><mi>Σ</mi><mi>x</mi><mo>−</mo><mi>b</mi><mo>⊘</mo><mi>x</mi></mrow></math>""",
        6,
    )
    _add_equation(
        doc,
        equations,
        """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msup><mo>∇</mo><mn>2</mn></msup><mi>f</mi><mo stretchy="false">（</mo><mi>x</mi><mo stretchy="false">）</mo><mo>=</mo><mi>Σ</mi><mo>+</mo><mi>diag</mi><mo stretchy="false">（</mo><msub><mi>b</mi><mi>i</mi></msub><mo>/</mo><msubsup><mi>x</mi><mi>i</mi><mn>2</mn></msubsup><mo stretchy="false">）</mo><mo>≻</mo><mn>0</mn></mrow></math>""",
        7,
    )
    _add_equation(
        doc,
        equations,
        """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mi>Σ</mi><mi>x</mi><mo>−</mo><mi>b</mi><mo>⊘</mo><mi>x</mi><mo>=</mo><mn>0</mn><mspace width="0.5em"/><mo>⇒</mo><mspace width="0.5em"/><msub><mi>x</mi><mi>i</mi></msub><msub><mrow><mo stretchy="false">（</mo><mi>Σ</mi><mi>x</mi><mo stretchy="false">）</mo></mrow><mi>i</mi></msub><mo>=</mo><msub><mi>b</mi><mi>i</mi></msub></mrow></math>""",
        8,
    )

    # 四、历史演进
    _add_heading(doc, 1, "四、历史优化器演进与受控复现")
    _add_heading(doc, 2, "（一）从原始 SLSQP 到数值尺度修补")
    _add_body(
        doc,
        "v0.01 使用原始风险贡献平方差和 SLSQP。v0.02 修正了收益率百分数单位及资产名称，但优化目标和算法设置没有变化，"
        "因此两版历史收益差异混合了数据修正影响，不能用于证明优化器进步。受控实验以 v0.02 作为原始目标基线，"
        "在默认停止条件下，SLSQP 往往只迭代一次便返回 success=True；由于目标值极小，通用停止准则误把初始等权点视为已收敛。"
    )
    _add_body(
        doc,
        "v0.03 将原目标严格乘以 10⁹，并把 ftol 收紧至 10⁻¹⁰、最大迭代次数提高至 1000。目标放大不改变理论最优解，"
        "但放大了函数差与梯度，使 SLSQP 更容易识别下降方向。v0.04 改用相对风险贡献误差，并在组合方差很低时设置保护项；"
        "该目标对协方差的正比例缩放保持不变，属于比固定倍数更有解释力的无量纲尺度处理。两版改动缓解了提前停止，"
        "但仍然依赖非凸误差目标和通用约束求解器。"
    )
    _add_heading(doc, 2, "（二）v0.05 的理论跃迁与 v0.06 的风险矩阵修正")
    _add_body(
        doc,
        "v0.05 采用式（5）的对数障碍凸模型，并为 L-BFGS-B 提供解析梯度。这不是简单的参数调优，而是把非凸风险贡献误差"
        "转化为严格凸无约束正域问题：最优点唯一，梯度具有清晰 KKT 含义，解后再归一化即可满足权重和约束。v0.06 引入 EWMA"
        "下行二阶矩和 10⁻⁸I 正则化，改善风险矩阵的经济口径与数值条件。后者属于输入风险模型升级，必须与求解器结构变化分开叙述。"
    )
    _add_heading(doc, 2, "（三）统一口径实验结果")
    optimizer_rows = []
    for _, row in optimizer.iterrows():
        optimizer_rows.append(
            (
                row["label"],
                row["objective_family"],
                _pct(row["solver_success_rate"], 1),
                _pct(row["rc_pass_rate"], 1),
                _num(row["median_iterations"], 1),
                _sci(row["median_rc_error"]),
            )
        )
    optimizer_table = _add_table(
        doc,
        state,
        "历史优化器受控复现实验汇总",
        ["版本", "目标/算法", "求解器成功率", "RC验收率", "中位迭代", "中位RC误差"],
        optimizer_rows,
        [1.25, 1.35, 0.9, 0.8, 0.75, 1.13],
        font_size=8.6,
    )
    optimizer_figure = _add_figure(
        doc, state, figures_dir / "optimizer_evolution.png", "历史优化器的 RC 验收率与中位误差", width=6.05
    )
    _add_body_with_ref(
        doc,
        "表 ",
        optimizer_table,
        " 给出了 148 个滚动矩阵上的统一汇总，图 ",
        str(state.table),
    )
    _add_body_with_ref(
        doc,
        "",
        optimizer_figure,
        " 进一步展示验收率与误差量级。v0.02 的求解器成功率虽然为 100%，RC 验收率却为 0，直接说明 success 标志不足以判断"
        "投资组合是否实现风险预算。目标放大和相对误差显著改善结果，但仍存在未达到 10⁻⁶ 阈值的窗口；凸重构后误差进一步下降，"
        "当前阻尼牛顿法在全部窗口通过独立验收。",
        str(state.figure),
    )
    _add_body(
        doc,
        "该比较只支持数值求解层面的结论，不支持把不同历史版本的回测收益差异归因于优化器。历史代码同时改变过收益单位、"
        "风险估计、资产可用集合与组合构建规则；如果直接比较版本净值，优化器贡献将与数据和策略变化混杂。受控复现的意义"
        "正是隔离这些干扰，仅考察同一数学问题下不同目标和算法设置能否得到合格解。"
    )

    # 五、牛顿法
    _add_heading(doc, 1, "五、阻尼牛顿法与数值性能")
    _add_heading(doc, 2, "（一）Newton 方向、正权重步长与 Armijo 回溯")
    _add_body(
        doc,
        "当前求解器在每次迭代中通过线性方程 Hₖpₖ=−gₖ 计算 Newton 方向，而不显式求逆。由于 Hessian 正定，"
        "只要梯度非零便有 gₖᵀpₖ<0，Newton 方向是严格下降方向。纯 Newton 步可能使某些 xᵢ 变为非正，"
        "因此先根据负方向分量计算最大可行步长，再乘以安全系数；随后执行 Armijo 回溯，直到目标函数满足充分下降条件。"
    )
    _add_equation(
        doc,
        equations,
        """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msub><mi>H</mi><mi>k</mi></msub><msub><mi>p</mi><mi>k</mi></msub><mo>=</mo><mo>−</mo><msub><mi>g</mi><mi>k</mi></msub><mo>,</mo><mspace width="0.5em"/><msub><mi>x</mi><mrow><mi>k</mi><mo>+</mo><mn>1</mn></mrow></msub><mo>=</mo><msub><mi>x</mi><mi>k</mi></msub><mo>+</mo><msub><mi>α</mi><mi>k</mi></msub><msub><mi>p</mi><mi>k</mi></msub></mrow></math>""",
        9,
    )
    _add_equation(
        doc,
        equations,
        """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mi>f</mi><mo stretchy="false">（</mo><mi>x</mi><mo>+</mo><mi>α</mi><mi>p</mi><mo stretchy="false">）</mo><mo>≤</mo><mi>f</mi><mo stretchy="false">（</mo><mi>x</mi><mo stretchy="false">）</mo><mo>+</mo><mi>c</mi><mi>α</mi><msup><mi>g</mi><mi>T</mi></msup><mi>p</mi></mrow></math>""",
        10,
    )
    _add_body(doc, "算法流程可以概括为以下五步：", first_indent=False)
    for item in (
        "以正向量作为初始点，计算目标、解析梯度和解析 Hessian。",
        "求解正定线性方程得到 Newton 方向，并检查下降性与数值有限性。",
        "根据 pᵢ<0 的分量限制最大步长，保证更新后的所有 xᵢ 仍为正。",
        "执行 Armijo 回溯，直到目标函数达到充分下降。",
        "同时检查梯度无穷范数、步长、风险贡献误差和最大迭代次数；归一化后再次验证权重和与 RC 误差。",
    ):
        _add_list_item(doc, item)
    _add_heading(doc, 2, "（二）与 L-BFGS-B、SLSQP 的比较")
    algorithm_rows = []
    for _, row in algorithms.sort_values("median_iterations").iterrows():
        algorithm_rows.append(
            (
                _method_label(row["method"]),
                _pct(row["success_rate"], 1),
                _num(row["median_iterations"], 1),
                _num(row["median_runtime_ms"], 2),
                _sci(row["median_rc_error"]),
                _sci(row["max_rc_error"]),
            )
        )
    algorithm_table = _add_table(
        doc,
        state,
        "当前三类求解器的统一比较",
        ["方法", "成功率", "中位迭代", "中位耗时/ms", "中位RC误差", "最大RC误差"],
        algorithm_rows,
        [1.05, 0.8, 0.8, 1.0, 1.2, 1.33],
        font_size=8.8,
    )
    convergence_figure = _add_figure(
        doc, state, figures_dir / "solver_convergence.png", "代表窗口的风险贡献误差收敛过程", width=5.9
    )
    solver_figure = _add_figure(
        doc, state, figures_dir / "solver_summary.png", "三类求解器的迭代次数、耗时与误差比较", width=6.0
    )
    _add_body_with_ref(
        doc,
        "表 ",
        algorithm_table,
        f" 显示，阻尼牛顿法中位迭代为 {newton['median_iterations']:.0f} 次，中位 RC 误差为 {_sci(newton['median_rc_error'])}；"
        "L-BFGS-B 同样利用解析梯度但不显式使用 Hessian，迭代次数较多；SLSQP 在直接 RC 误差目标上还需同时处理权重和约束，"
        "精度和稳定性更依赖停止参数。这里的运行时间只反映本机环境，主要结论应以误差和迭代行为为准。",
        str(state.table),
    )
    _add_body_with_ref(
        doc,
        "代表窗口的逐次误差见图 ",
        convergence_figure,
        "，汇总比较见图 ",
        str(state.figure - 1),
    )
    _add_body_with_ref(
        doc,
        "",
        solver_figure,
        "。解析 Hessian 使 Newton 法在接近最优点时表现出快速局部收敛；阻尼和回溯则补足了远离最优点时的全局化机制。"
        "当前实现并未把“迭代少”作为唯一成功标准，而是要求最终风险贡献和约束同时通过独立诊断。",
        str(state.figure),
    )
    _add_heading(doc, 2, "（三）病态风险矩阵压力测试")
    stress_1e8 = stress.loc[stress["condition_number"] == stress["condition_number"].max()]
    stress_rows = []
    for _, row in stress_1e8.iterrows():
        stress_rows.append(
            (
                _method_label(row["method"]),
                _pct(row["success_rate"], 1),
                _num(row["median_iterations"], 1),
                _sci(row["median_rc_error"]),
                _sci(row["max_rc_error"]),
            )
        )
    stress_table = _add_table(
        doc,
        state,
        f"条件数 {stress['condition_number'].max():.0e} 下的压力测试",
        ["方法", "成功率", "中位迭代", "中位RC误差", "最大RC误差"],
        stress_rows,
        [1.1, 0.9, 0.95, 1.55, 1.68],
        font_size=9.2,
    )
    stress_figure = _add_figure(doc, state, figures_dir / "stress_test.png", "病态矩阵下的风险贡献误差", width=5.95)
    _add_body_with_ref(
        doc,
        "表 ",
        stress_table,
        " 和图 ",
        str(state.table),
    )
    _add_body_with_ref(
        doc,
        "",
        stress_figure,
        " 展示条件数逐步提高时的数值表现。正则化并不能消除所有估计风险，但能使风险矩阵保持正定；"
        "阻尼牛顿法在条件数 10⁸ 的压力场景中仍维持较低 RC 误差。压力测试说明二阶结构和解后验收对病态输入尤其重要，"
        "但其结论仅限于构造矩阵上的数值稳定性，不等同于真实市场中的收益稳健性。",
        str(state.figure),
    )

    # 六、实证
    _add_heading(doc, 1, "六、样本外回测与稳健性分析")
    _add_heading(doc, 2, "（一）策略设置与净值路径")
    _add_body(
        doc,
        "实证部分比较等权、逆下行波动率和 ERC 风险平价三类策略。等权组合作为不使用风险估计的基准；逆下行波动率仅根据"
        "单资产下行波动配置资金，不考虑相关性；ERC 同时使用协方差结构和风险预算。三类策略采用相同调仓日和成本口径，"
        "从而把差异集中在权重生成规则。"
    )
    nav_figure = _add_figure(doc, state, figures_dir / "strategy_nav.png", "三类策略的累计净值", width=6.05)
    _add_body_with_ref(
        doc,
        "累计净值路径见图 ",
        nav_figure,
        "。等权组合在部分权益上涨阶段获得更高绝对收益，但其风险暴露更集中于高波动资产；风险平价净值路径相对平滑。"
        "判断风险平价是否有效不能只看终点净值，还应同时考察波动率、最大回撤、夏普比率和换手成本。",
        str(state.figure),
    )
    _add_heading(doc, 2, "（二）验证期风险收益表现")
    validation_rows = []
    for _, row in strategies.loc[strategies["period"] == "validation"].iterrows():
        validation_rows.append(
            (
                _strategy_label(row["strategy"]),
                _pct(row["annual_return"]),
                _pct(row["annual_volatility"]),
                _num(row["sharpe"]),
                _pct(row["max_drawdown"]),
                _num(row["annual_turnover"]),
            )
        )
    validation_table = _add_table(
        doc,
        state,
        "验证期策略绩效",
        ["策略", "年化收益", "年化波动", "夏普比率", "最大回撤", "年换手"],
        validation_rows,
        [1.28, 0.9, 0.9, 0.8, 0.9, 1.4],
        font_size=9.2,
    )
    _add_body_with_ref(
        doc,
        "验证期结果见表 ",
        validation_table,
        f"。风险平价年化收益为 {_pct(validation_erc['annual_return'])}，低于等权组合的 "
        f"{_pct(validation_equal['annual_return'])}；但其年化波动率由 {_pct(validation_equal['annual_volatility'])} "
        f"降至 {_pct(validation_erc['annual_volatility'])}，最大回撤由 {_pct(validation_equal['max_drawdown'])} "
        f"收窄至 {_pct(validation_erc['max_drawdown'])}，夏普比率提高至 {_num(validation_erc['sharpe'])}。"
        "因此，核心实证结论是风险调整后表现和回撤控制改善，而不是风险平价获得最高绝对收益。",
        str(state.table),
    )
    _add_heading(doc, 2, "（三）风险估计口径与年度情景")
    estimator_rows = []
    for _, row in estimators.loc[estimators["period"] == "validation"].iterrows():
        estimator_rows.append(
            (
                _estimator_label(row["estimator"]),
                _pct(row["annual_return"]),
                _pct(row["annual_volatility"]),
                _num(row["sharpe"]),
                _pct(row["max_drawdown"]),
            )
        )
    estimator_table = _add_table(
        doc,
        state,
        "验证期风险估计方法比较",
        ["风险矩阵", "年化收益", "年化波动", "夏普比率", "最大回撤"],
        estimator_rows,
        [1.75, 1.0, 1.0, 0.95, 1.48],
        font_size=9.3,
    )
    estimator_figure = _add_figure(
        doc, state, figures_dir / "estimator_comparison.png", "不同风险估计方法的验证期表现", width=5.95
    )
    yearly_figure = _add_figure(doc, state, figures_dir / "yearly_returns.png", "三类策略的年度收益与回撤情景", width=6.0)
    _add_body_with_ref(
        doc,
        "表 ",
        estimator_table,
        " 与图 ",
        str(state.table),
    )
    _add_body_with_ref(
        doc,
        "",
        estimator_figure,
        " 表明，样本协方差和 EWMA 全协方差在当前验证期的夏普比率高于主模型的 EWMA 半协方差。"
        "因此，半协方差并非由结果倒推得到的“最优估计器”，而是基于更关注下行风险的建模选择。年度情景见图 ",
        str(state.figure - 1),
    )
    _add_body_with_ref(
        doc,
        "",
        yearly_figure,
        "。2026 年仅包含截至 4 月初的数据，不能与完整年度直接比较。不同年份的收益排序并不稳定，"
        "再次说明实证结论应聚焦长期风险结构而非单年胜负。",
        str(state.figure),
    )
    _add_heading(doc, 2, "（四）训练期选参与参数敏感性")
    sensitivity_figure = _add_figure(
        doc, state, figures_dir / "sensitivity_heatmap.png", "窗口长度和 EWMA 衰减系数的训练期敏感性", width=5.9
    )
    _add_body_with_ref(
        doc,
        "参数网格结果见图 ",
        sensitivity_figure,
        f"。按训练期夏普比率最高、再以最大回撤较轻作为并列规则，选中窗口 {selected['window']}、"
        f"衰减系数 {selected['decay']:.2f}；其训练期夏普为 {_num(selected['train_sharpe'])}，验证期夏普为 "
        f"{_num(selected['validation_sharpe'])}。主报告仍保留项目既定的 252 日、0.97 参数作为核心模型，以维持历史口径；"
        "网格结果用于说明结论是否依赖单点设置，而不是在验证期上重新调参。",
        str(state.figure),
    )
    _add_body(
        doc,
        "敏感性表覆盖多个窗口与衰减系数组合，大部分验证期风险平价结果仍表现为较低波动和较小回撤，但夏普差异说明风险估计"
        "参数会改变权重反应速度与换手。较短窗口更快追随市场状态，也可能放大估计噪声；较高衰减系数使用更长记忆，"
        "但对结构突变响应较慢。参数选择应在稳定性、及时性和交易成本之间权衡。"
    )

    # 七、代表解
    _add_heading(doc, 1, "七、代表窗口的最优权重与风险贡献")
    _add_heading(doc, 2, "（一）资金权重与等风险贡献")
    representative_rows = []
    for _, row in representative.iterrows():
        representative_rows.append(
            (
                row["asset"],
                _pct(row["weight"]),
                _pct(row["risk_contribution"]),
                _pct(row["target_risk_contribution"]),
                _sci(abs(row["risk_contribution"] - row["target_risk_contribution"])),
            )
        )
    representative_table = _add_table(
        doc,
        state,
        f"{representative['representative_date'].iloc[0]} 代表窗口的 ERC 解",
        ["资产", "资金权重", "实际RC", "目标RC", "绝对误差"],
        representative_rows,
        [1.45, 1.0, 1.0, 1.0, 1.73],
        font_size=9.1,
    )
    weight_figure = _add_figure(
        doc, state, figures_dir / "weights_risk_contributions.png", "代表窗口的资金权重与风险贡献", width=6.05
    )
    _add_body_with_ref(
        doc,
        "代表窗口结果见表 ",
        representative_table,
        " 和图 ",
        str(state.table),
    )
    _add_body_with_ref(
        doc,
        "",
        weight_figure,
        "。各资产资金权重明显不同，但实际风险贡献均接近 11.11%。债券或低波动资产需要较高资金权重，"
        "高波动权益和商品资产则以较低资金权重贡献相同风险。该结果直观说明风险平价不是等权配置，也不是简单的逆波动率配置；"
        "相关性通过 Σw 项进入边际风险贡献，对最优权重具有实质影响。",
        str(state.figure),
    )
    _add_heading(doc, 2, "（二）解的诊断意义")
    _add_body(
        doc,
        "当前求解器在归一化后重新计算风险贡献，并同时记录最大 RC 误差、权重和误差与最小权重。即使优化目标或梯度满足停止条件，"
        "若风险贡献偏差超过 10⁻⁶，结果仍不通过验收。这一设计把数学停止条件与金融目标检查分开：前者判断数值迭代是否结束，"
        "后者判断所得权重是否真正实现风险预算。"
    )

    # 八、扩展
    _add_heading(doc, 1, "八、动态可行域、一般风险预算与外部约束")
    _add_heading(doc, 2, "（一）动态可行资产集合")
    _add_body(
        doc,
        "v0.15 按照资产上市状态改变进入优化器的资产集合。若某资产在当期不可交易，则对应权重不应出现在决策变量中；"
        "模型需要在随时间变化的可行域上重新归一化风险预算。这属于约束集合随时间变化，而不是目标函数或 Newton 方向本身的变化。"
        "本文主回测使用项目已提供的连续指数代理以维持资产类别历史，但在解释实际可交易性时明确披露这一限制。"
    )
    _add_heading(doc, 2, "（二）一般风险预算")
    _add_body(
        doc,
        "v0.16_2 将 bᵢ=1/n 推广为信号条件化的正预算向量。风险预算决定资产应承担的风险份额，而不直接指定资金权重；"
        "当协方差为对角矩阵时，闭式解满足 wᵢ∝√(bᵢ/Σᵢᵢ)，说明预算提高与波动率共同决定资金配置。该闭式关系已纳入单元测试。"
    )
    rb_math = """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msub><mi>w</mi><mi>i</mi></msub><mo>∝</mo><msqrt><mfrac><msub><mi>b</mi><mi>i</mi></msub><msub><mi>Σ</mi><mrow><mi>i</mi><mi>i</mi></mrow></msub></mfrac></msqrt><mspace width="0.5em"/><mo stretchy="false">（</mo><mi>Σ</mi><mspace width="0.2em"/><mi>为对角矩阵</mi><mo stretchy="false">）</mo></mrow></math>"""
    _add_equation(doc, equations, rb_math, 11)
    rb_rows = []
    for _, row in risk_budget.iterrows():
        rb_rows.append(
            (
                row["asset"],
                _num(row["raw_budget_multiplier"], 1),
                _pct(row["target_risk_budget"]),
                _pct(row["actual_risk_contribution"]),
                _pct(row["weight"]),
                _sci(row["absolute_rc_error"]),
            )
        )
    rb_table = _add_table(
        doc,
        state,
        f"{rb_summary['representative_date']} 一般风险预算代表解",
        ["资产", "预算倍率", "目标预算", "实际RC", "资金权重", "RC误差"],
        rb_rows,
        [1.32, 0.82, 0.9, 0.9, 0.9, 1.34],
        font_size=8.8,
    )
    rb_figure = _add_figure(
        doc, state, figures_dir / "risk_budget_extension.png", "一般风险预算的目标贡献、实际贡献与资金权重", width=6.05
    )
    _add_body_with_ref(
        doc,
        "代表实验将沪深 300ETF 和中证 1000ETF 的原始预算倍率设为 2，其余资产设为 1。归一化后，两类权益资产的目标风险预算"
        "分别为 18.18%，其余资产为 9.09%。结果见表 ",
        rb_table,
        " 和图 ",
        str(state.table),
    )
    _add_body_with_ref(
        doc,
        "",
        rb_figure,
        f"。阻尼牛顿法的最大预算跟踪误差为 {_sci(rb_summary['rc_max_error'])}，说明同一凸模型能够准确实现非等额风险预算。"
        "该实验只验证模型推广能力，并不证明提高权益风险预算会改善未来收益。",
        str(state.figure),
    )
    _add_heading(doc, 2, "（三）外部波动率覆盖")
    _add_body(
        doc,
        "v0.16_3 先根据近 60 日实现波动率缩放股指资金仓位，再将剩余资金交给风险预算模块。该规则在凸优化器之外改变了组合资金覆盖，"
        "既没有修改式（5）的目标，也没有修改梯度或 Hessian。因此，它应被描述为组合构建层的外部风险覆盖，而不是优化器升级。"
        "这一层级区分对于解释历史策略迭代十分重要：模型、求解器、风险估计、可行域和后处理可以同时演进，但各自解决的问题不同。"
    )

    # 九、结论
    _add_heading(doc, 1, "九、结论与局限")
    _add_heading(doc, 2, "（一）主要结论")
    _add_body(
        doc,
        "第一，历史迭代表明，目标放大和相对误差能够缓解尺度敏感与提前停止，但真正的理论跃迁来自对数障碍凸重构。"
        "严格凸性保证正域内最优解唯一，KKT 条件把最优点直接映射为目标风险预算。第二，当前阻尼牛顿法利用解析 Hessian、"
        "正权重步长与 Armijo 回溯，在统一滚动风险矩阵上以较少迭代达到高精度，并通过独立 RC 验收和病态矩阵压力测试。"
    )
    _add_body(
        doc,
        "第三，样本外结果显示风险平价的优势主要体现在较低波动率、较小最大回撤和较高风险调整后收益，而非最高绝对收益。"
        "第四，一般风险预算可以在不直接指定资金权重的情况下表达非等额风险偏好；动态可行域和外部波动率覆盖则属于更上层的"
        "资产可用性与组合构建规则。对历史版本的评价必须保持这些层级边界。"
    )
    _add_heading(doc, 2, "（二）研究局限")
    _add_body(
        doc,
        "本研究仍存在四方面局限。其一，ETF 成立前使用指数代理，真实可交易成本、跟踪误差和流动性并未完全进入模型。其二，"
        "风险矩阵仅比较样本协方差、EWMA 全协方差和 EWMA 半协方差，尚未纳入更系统的收缩估计、因子模型或状态切换模型。"
        "其三，回测采用固定 5bp 单边成本和月度调仓，未刻画冲击成本、申赎约束与成交容量。其四，参数网格有限，2026 年又是不完整年度，"
        "因此样本外结果不能外推为未来收益保证。"
    )
    _add_heading(doc, 2, "（三）后续研究方向")
    _add_body(
        doc,
        "后续可在保持风险预算凸结构的基础上，引入协方差收缩、权重上下限、换手惩罚和滚动交叉验证；对于带线性约束的一般风险预算，"
        "可进一步研究障碍法、坐标下降或锥优化实现。还可把算法误差与风险估计误差分解，比较“更精确求解一个估计模型”"
        "与“改进风险矩阵估计”对最终组合的相对重要性。"
    )

    # 参考文献
    _add_heading(doc, 1, "参考文献")
    references = [
        "[1] MARKOWITZ H. Portfolio selection[J]. The Journal of Finance, 1952, 7(1): 77-91.",
        "[2] MAILLARD S, RONCALLI T, TEÏLETCHE J. The properties of equally weighted risk contribution portfolios[J]. The Journal of Portfolio Management, 2010, 36(4): 60-70.",
        "[3] BRUDER B, RONCALLI T. Managing risk exposures using the risk budgeting approach[R]. SSRN Working Paper 2009778, 2012.",
        "[4] RONCALLI T. Introduction to Risk Parity and Budgeting[M]. Boca Raton: CRC Press, 2013.",
        "[5] SPINU F. An algorithm for computing risk parity weights[R]. SSRN Working Paper 2297383, 2013, revised 2020.",
        "[6] BOYD S, VANDENBERGHE L. Convex Optimization[M]. Cambridge: Cambridge University Press, 2004.",
        "[7] NOCEDAL J, WRIGHT S J. Numerical Optimization[M]. 2nd ed. New York: Springer, 2006.",
        "[8] J.P. MORGAN, REUTERS. RiskMetrics Technical Document[R]. 4th ed. New York: Morgan Guaranty Trust Company, 1996.",
        "[9] QIAN E. Risk parity portfolios: Efficient portfolios through true diversification[R]. PanAgora Asset Management, 2005.",
        "[10] 华泰研究. 从资产配置走向因子配置：中国版全天候增强策略[R]. 2025.",
        "[11] 本项目历史版本代码. v0.01—v0.06、v0.15、v0.16_2、v0.16_3[CP].",
    ]
    for reference in references:
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.paragraph_format.first_line_indent = Pt(-21)
        paragraph.paragraph_format.left_indent = Pt(21)
        paragraph.paragraph_format.line_spacing = 1.2
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(reference)
        _set_run_font(run, size=10.5)

    # 附录
    _add_heading(doc, 1, "附录 A　复现与验收口径")
    _add_body(
        doc,
        "运行项目根目录下的 run.py 可依次重建数值实验、策略回测、静态图表、交互式 HTML、Word 报告和 Word PDF。"
        "Word、PDF、HTML 与 CSV/JSON 共用同一结果源，报告构建阶段不重新估计模型或手工抄写数值。"
    )
    appendix_rows = [
        ("历史优化器记录", "optimizer_evolution_details.csv", "148×5=740 条诊断"),
        ("当前算法记录", "algorithm_monthly_details.csv", "三类求解器逐月诊断"),
        ("回测结果", "strategy_metrics.csv / strategy_nav.csv", "训练期、验证期与全样本"),
        ("一般风险预算", "risk_budget_extension.csv", "目标预算、实际RC和资金权重"),
        ("验证清单", "pdf_verification.json", "Word引擎、页数、术语、哈希和结构"),
    ]
    _add_table(
        doc,
        state,
        "主要复现文件与验收用途",
        ["内容", "文件", "验收用途"],
        appendix_rows,
        [1.35, 2.55, 2.28],
        font_size=9.3,
    )
    _add_body(
        doc,
        f"敏感性实验共包含 {len(sensitivity)} 条参数—期间组合记录。最终验收不以 SciPy success 标志代替金融目标检查；"
        "所有求解器均要求权重非负、权重和接近 1，并以最大绝对风险贡献误差不超过 10⁻⁶ 作为 RC 通过标准。"
    )
    _add_heading(doc, 2, "（一）统一验收指标")
    _add_body(
        doc,
        "每个滚动窗口同时保留求解器状态、迭代次数、耗时、RC 误差、权重和误差与最小权重。求解器状态回答程序是否按内部规则结束，"
        "RC 与约束误差回答所得解是否满足金融模型；受控复现统一采用 10⁻⁶ 阈值，两类判断互不替代。"
    )
    _add_heading(doc, 2, "（二）可审计输出与数字一致性")
    _add_body(
        doc,
        "optimizer_evolution_details.csv 应包含 148 个日期与 5 个历史变体，共 740 条诊断；汇总表只由该明细表聚合。"
        "策略绩效、代表窗口权重和一般风险预算分别回溯到对应 CSV；构建程序检查 DOCX/PDF 的关键数值和术语，"
        "并把文件哈希、Word 版本、页数与域计数写入验证 JSON。"
    )
    _add_heading(doc, 2, "（三）环境与可迁移性说明")
    _add_body(
        doc,
        "耗时结果依赖处理器、Python、SciPy 和底层线性代数库，只能作为同一机器上的相对参考；论文结论主要依据 RC 验收率、"
        "误差分布和迭代行为。DOCX 保留自动目录、页码、题注、交叉引用与原生公式，便于在 Microsoft Word 中继续编辑；"
        "若使用其他办公软件打开，应重新检查分页，但不得据其分页替代 Word 导出的提交版 PDF。"
    )


def _add_report_content(doc: DocumentObject, course_dir: Path, config: dict[str, Any],
                        summary: dict[str, Any]) -> None:
    """Build the compact, theory-led submission report."""

    tables_dir = course_dir / "output" / "tables"
    figures_dir = course_dir / "output" / "figures"
    data_quality = summary["data_quality"]
    validation_erc = summary["validation_erc"]
    validation_equal = summary["validation_equal_weight"]
    algorithms = _load_frame(tables_dir, "algorithm_summary.csv")
    stress = _load_frame(tables_dir, "stress_test_summary.csv")
    strategies = _load_frame(tables_dir, "strategy_metrics.csv")
    representative = _load_frame(tables_dir, "representative_optimal_solution.csv")
    risk_budget = _load_frame(tables_dir, "risk_budget_extension.csv")
    selected = summary["selected_parameter"]
    rb_summary = summary["risk_budget_extension"]
    newton = summary["newton_summary"]
    state = FieldState()
    equations = EquationTransformer()

    # 一、引言
    _add_heading(doc, 1, "一、引言")
    _add_heading(doc, 2, "（一）研究背景")
    _add_body(
        doc,
        "Markowitz 均值—方差模型把资产配置形式化为收益与风险之间的优化问题，但最优权重对预期收益和协方差输入较为敏感，"
        "有限样本中的微小估计误差可能造成明显的仓位变化[1]。当投资范围同时包含股票、债券、商品和黄金时，资金等权也不等于"
        "风险分散：高波动或与其他资产高度相关的品种可能承担组合中的大部分风险。风险预算方法因此把决策视角从资本比例转向"
        "风险贡献比例，使投资者能够在不依赖精确收益预测的条件下表达分散化目标。"
    )
    _add_body(
        doc,
        "风险贡献相等并不要求资金权重相等。低波动、低相关资产往往需要较高资金权重，才能与高波动资产承担相近风险。"
        "这一机制具有清晰的金融含义，却对应一个正权重域上的非线性方程组；若直接最小化风险贡献残差，目标函数还会受到"
        "风险矩阵尺度和通用停止条件影响。因此，本报告把重点放在风险预算问题的凸等价重构、二阶数值求解和独立结果验收上。"
    )
    _add_heading(doc, 2, "（二）相关研究")
    _add_body(
        doc,
        "Tasche 依据一阶齐次风险度量和 Euler 定理说明，总风险可以分解为各仓位与其边际风险的乘积之和，且该分解具有完整分配"
        "和风险调整绩效相容的经济含义[2]。Maillard、Roncalli 与 Teïletche 系统研究了等风险贡献组合，指出其波动率通常位于"
        "最小方差组合与资金等权组合之间，在风险集中和权重集中之间提供折中[3]。这些结论为 ERC 的风险分散解释提供了理论基础。"
    )
    _add_body(
        doc,
        "Bruder 与 Roncalli 将 ERC 推广为任意正预算下的一般风险预算组合，强调风险预算描述的是资产应承担的总风险份额，而不是"
        "预先指定资金仓位[4]。Cetingoz、Fermanian 与 Guéant 进一步从凸优化角度讨论风险预算组合的存在性、唯一性及计算方法[5]。"
        "在数值算法方面，Spinu 将风险预算方程转化为可由 Newton 类方法高效求解的凸问题[6]；阻尼 Newton 方向和线搜索的"
        "一般理论可参见 Nocedal 与 Wright[7]。上述云盘文献构成本文第一章的主要理论脉络。"
    )
    _add_body(
        doc,
        "综合上述研究，风险平价并不是脱离均值—方差框架的另一套风险定义，而是对协方差风险进行不同的配置：传统模型由预期收益"
        "驱动最优权重，风险预算模型则直接规定各资产对总风险的承担比例。前者更适合表达明确的收益观点，后者在收益预测不稳定时"
        "更强调可解释的风险分散。本文据此不把ERC视为收益最大化策略，而把研究重点限定为风险预算条件的可解性、算法精度和样本外"
        "风险特征，避免用单一累计收益替代理论与数值评价。"
    )
    _add_heading(doc, 2, "（三）研究问题与主要贡献")
    _add_body(
        doc,
        "本文回答三个问题：ERC 条件能否转化为具有唯一解的严格凸模型；如何利用解析 Hessian、正权重步长和 Armijo 回溯"
        "稳定求解；所得权重能否在滚动风险矩阵和样本外资产配置中满足风险贡献目标。相应贡献包括：给出风险预算模型及 KKT"
        "等价关系，构建带独立 RC 验收的阻尼牛顿求解器，并以真实多资产数据验证其数值性质、风险控制效果和一般预算推广能力。"
    )

    # 二、数据与实验设计
    _add_heading(doc, 1, "二、数据与实验设计")
    _add_heading(doc, 2, "（一）资产范围与数据处理")
    _add_body(
        doc,
        f"实验使用项目原始 Excel 工作簿第 {data_quality['sheet_index'] + 1} 个工作表，样本期为 "
        f"{data_quality['start_date']} 至 {data_quality['end_date']}，共 {data_quality['rows']} 个交易日和 "
        f"{data_quality['assets']} 类资产。资产池覆盖大盘权益、小盘权益、红利低波、长久期国债、有色、能源化工、豆粕和黄金等"
        "ETF或指数代理。ETF成立前沿用项目提供的指数代理，因此回测反映连续的资产类别暴露，并不代表所有日期均可交易同一只基金。"
    )
    data_table = _add_table(
        doc,
        state,
        "样本与数据质量概览",
        ["项目", "统计结果"],
        [
            ("样本区间", f"{data_quality['start_date']}—{data_quality['end_date']}"),
            ("交易日与资产", f"{data_quality['rows']} 日，{data_quality['assets']} 类资产"),
            ("重复日期", str(data_quality["duplicate_dates"])),
            ("缺失值处理", f"{data_quality['missing_cells_before_fill']} 个孤立缺失收益填 0"),
            ("收益率范围", f"{_pct(data_quality['min_return'])} 至 {_pct(data_quality['max_return'])}"),
        ],
        [2.0, 4.18],
        font_size=10.0,
    )
    _add_body_with_ref(
        doc,
        "样本质量概览见表 ",
        data_table,
        "。原表收益率以百分数存储，进入模型前统一转换为小数；数据日期保持递增且不存在重复记录。",
        str(state.table),
    )
    _add_heading(doc, 2, "（二）EWMA 半协方差与正则化")
    _add_body(
        doc,
        "主模型使用 252 个交易日滚动窗口和 0.97 的 EWMA 衰减系数，较近观测获得更高权重；收益只保留负向部分，以强调亏损状态"
        "下的共同波动。为降低短样本和高度相关资产造成的近奇异问题，在风险矩阵对角线上加入 10⁻⁸I。该正则项用于改善线性方程"
        "求解条件，并不改变风险预算的经济定义。"
    )
    _add_equation(
        doc,
        equations,
        f"""<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mi>Σ</mi><mo>=</mo><munderover><mo>∑</mo><mrow><mi>t</mi><mo>=</mo><mn>1</mn></mrow><mi>T</mi></munderover><msub><mi>α</mi><mi>t</mi></msub><msub><mi>r</mi><mi>t</mi></msub><msubsup><mi>r</mi><mi>t</mi><mi>T</mi></msubsup><mo>+</mo><msup><mn>10</mn><mrow><mo>−</mo><mn>8</mn></mrow></msup><mi>I</mi><mo>,</mo><mspace width="0.5em"/><msub><mi>α</mi><mi>t</mi></msub><mo>∝</mo><msup><mn>{config['decay']}</mn><mrow><mi>T</mi><mo>−</mo><mi>t</mi></mrow></msup></mrow></math>""",
        1,
    )
    _add_heading(doc, 2, "（三）滚动回测与评价指标")
    _add_body(
        doc,
        f"每个月末使用截至当日的 {config['window']} 日数据估计风险矩阵，目标权重在下一交易日执行，单边换手成本为 "
        f"{config['fee_rate'] * 10000:.0f}bp。训练期为 {config['train_start']} 至 {config['train_end']}，验证期为 "
        f"{config['validation_start']} 至 {config['validation_end']}，参数选择仅使用训练期。数值评价记录求解状态、迭代次数、"
        "RC误差和权重约束误差；实证评价使用年化收益、年化波动、夏普比率、最大回撤和换手率。"
    )

    # 三、ERC模型与凸等价重构
    _add_heading(doc, 1, "三、ERC模型与凸等价重构")
    _add_heading(doc, 2, "（一）组合风险与 Euler 分解")
    _add_body(
        doc,
        "设资金权重为 w，协方差矩阵为 Σ，且 wᵢ≥0、1ᵀw=1。组合波动率是一阶齐次函数，其边际风险贡献为波动率对各权重的"
        "偏导；根据 Euler 定理，权重与边际贡献之积可以完全加总为组合波动率[2]。将各资产总风险贡献除以组合总风险，得到"
        "归一化风险贡献 RCᵢ。"
    )
    _add_equation(doc, equations, """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msub><mi>σ</mi><mi>p</mi></msub><mo stretchy="false">（</mo><mi>w</mi><mo stretchy="false">）</mo><mo>=</mo><msqrt><msup><mi>w</mi><mi>T</mi></msup><mi>Σ</mi><mi>w</mi></msqrt></mrow></math>""", 2)
    _add_equation(doc, equations, """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msub><mi>MRC</mi><mi>i</mi></msub><mo>=</mo><mfrac><msub><mrow><mo stretchy="false">（</mo><mi>Σ</mi><mi>w</mi><mo stretchy="false">）</mo></mrow><mi>i</mi></msub><msub><mi>σ</mi><mi>p</mi></msub></mfrac><mo>,</mo><mspace width="0.5em"/><msub><mi>σ</mi><mi>p</mi></msub><mo>=</mo><munderover><mo>∑</mo><mrow><mi>i</mi><mo>=</mo><mn>1</mn></mrow><mi>n</mi></munderover><msub><mi>w</mi><mi>i</mi></msub><msub><mi>MRC</mi><mi>i</mi></msub></mrow></math>""", 3)
    _add_equation(doc, equations, """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msub><mi>RC</mi><mi>i</mi></msub><mo>=</mo><mfrac><mrow><msub><mi>w</mi><mi>i</mi></msub><msub><mrow><mo stretchy="false">（</mo><mi>Σ</mi><mi>w</mi><mo stretchy="false">）</mo></mrow><mi>i</mi></msub></mrow><mrow><msup><mi>w</mi><mi>T</mi></msup><mi>Σ</mi><mi>w</mi></mrow></mfrac><mo>,</mo><mspace width="0.5em"/><munderover><mo>∑</mo><mrow><mi>i</mi><mo>=</mo><mn>1</mn></mrow><mi>n</mi></munderover><msub><mi>RC</mi><mi>i</mi></msub><mo>=</mo><mn>1</mn></mrow></math>""", 4)
    _add_heading(doc, 2, "（二）ERC条件与直接残差目标")
    _add_body(
        doc,
        "ERC要求每个资产承担 1/n 的组合风险；一般风险预算则要求 RCᵢ=bᵢ，其中 bᵢ>0 且 ∑bᵢ=1。一个直观做法是"
        "在非负和资金约束下最小化实际风险贡献与目标预算的平方差，但该残差目标通常不是全局凸函数，且函数值会随协方差"
        "正比例缩放而显著改变，通用停止阈值因而缺少统一尺度。"
    )
    _add_equation(doc, equations, """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msub><mi>RC</mi><mi>i</mi></msub><mo>=</mo><msub><mi>b</mi><mi>i</mi></msub><mo>,</mo><mspace width="0.4em"/><msub><mi>b</mi><mi>i</mi></msub><mo>&gt;</mo><mn>0</mn><mo>,</mo><mspace width="0.4em"/><munderover><mo>∑</mo><mrow><mi>i</mi><mo>=</mo><mn>1</mn></mrow><mi>n</mi></munderover><msub><mi>b</mi><mi>i</mi></msub><mo>=</mo><mn>1</mn></mrow></math>""", 5)
    _add_equation(doc, equations, """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><munderover><mo>min</mo><mrow><mi>w</mi><mo>≥</mo><mn>0</mn></mrow><mrow><mn>1</mn><mo>ᵀ</mo><mi>w</mi><mo>=</mo><mn>1</mn></mrow></munderover><mspace width="0.4em"/><munderover><mo>∑</mo><mrow><mi>i</mi><mo>=</mo><mn>1</mn></mrow><mi>n</mi></munderover><msup><mrow><mo stretchy="false">（</mo><msub><mi>RC</mi><mi>i</mi></msub><mo>−</mo><msub><mi>b</mi><mi>i</mi></msub><mo stretchy="false">）</mo></mrow><mn>2</mn></msup></mrow></math>""", 6)
    _add_heading(doc, 2, "（三）对数障碍凸模型")
    _add_body(
        doc,
        "在正权重域上，引入二次风险项与带预算系数的对数障碍项。若 Σ 正定且预算严格为正，二次项为凸函数，负对数项具有"
        "正定对角 Hessian，因而总目标严格凸并存在唯一极小点[5-6]。优化变量 x 不施加资金和约束；求得正向量后以"
        " w=x/(1ᵀx) 归一化即可恢复资金权重。"
    )
    _add_equation(doc, equations, """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><munderover><mo>min</mo><mi>x</mi><mrow><mi>x</mi><mo>&gt;</mo><mn>0</mn></mrow></munderover><mspace width="0.4em"/><mi>f</mi><mo stretchy="false">（</mo><mi>x</mi><mo stretchy="false">）</mo><mo>=</mo><mfrac><mn>1</mn><mn>2</mn></mfrac><msup><mi>x</mi><mi>T</mi></msup><mi>Σ</mi><mi>x</mi><mo>−</mo><munderover><mo>∑</mo><mrow><mi>i</mi><mo>=</mo><mn>1</mn></mrow><mi>n</mi></munderover><msub><mi>b</mi><mi>i</mi></msub><mi>log</mi><msub><mi>x</mi><mi>i</mi></msub></mrow></math>""", 7)
    _add_heading(doc, 2, "（四）梯度、Hessian与KKT条件")
    _add_body(
        doc,
        "解析梯度和Hessian既证明严格凸性，也直接用于Newton迭代。Hessian是正定风险矩阵与正对角矩阵之和，因此可行域内的"
        "Newton方向唯一。一阶最优条件逐元素乘以x后得到 xᵢ(Σx)ᵢ=bᵢ；再利用归一化和Euler分解，即得到 RCᵢ=bᵢ。"
        "因此，凸模型不是对风险预算方程的近似，而是其在正权重域中的等价求解形式。"
    )
    _add_body(
        doc,
        "该重构还区分了方向与尺度两个层次。风险贡献只取决于权重射线，最终资金约束通过归一化施加；凸目标中的对数项则为变量x"
        "选择唯一尺度。将一阶条件求和可得xᵀΣx=∑bᵢ=1，故最优点的尺度由预算总和自动固定。若经验风险矩阵仅为半正定，对角"
        "正则项使二次项具有严格曲率，并改善Hessian的条件数；因此正则化既服务于理论唯一性，也服务于有限精度下的线性求解。"
    )
    _add_equation(doc, equations, """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mo>∇</mo><mi>f</mi><mo stretchy="false">（</mo><mi>x</mi><mo stretchy="false">）</mo><mo>=</mo><mi>Σ</mi><mi>x</mi><mo>−</mo><mi>b</mi><mo>⊘</mo><mi>x</mi></mrow></math>""", 8)
    _add_equation(doc, equations, """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msup><mo>∇</mo><mn>2</mn></msup><mi>f</mi><mo stretchy="false">（</mo><mi>x</mi><mo stretchy="false">）</mo><mo>=</mo><mi>Σ</mi><mo>+</mo><mi>diag</mi><mo stretchy="false">（</mo><msub><mi>b</mi><mi>i</mi></msub><mo>/</mo><msubsup><mi>x</mi><mi>i</mi><mn>2</mn></msubsup><mo stretchy="false">）</mo><mo>≻</mo><mn>0</mn></mrow></math>""", 9)
    _add_equation(doc, equations, """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mi>Σ</mi><mi>x</mi><mo>−</mo><mi>b</mi><mo>⊘</mo><mi>x</mi><mo>=</mo><mn>0</mn><mspace width="0.5em"/><mo>⇒</mo><mspace width="0.5em"/><msub><mi>x</mi><mi>i</mi></msub><msub><mrow><mo stretchy="false">（</mo><mi>Σ</mi><mi>x</mi><mo stretchy="false">）</mo></mrow><mi>i</mi></msub><mo>=</mo><msub><mi>b</mi><mi>i</mi></msub></mrow></math>""", 10)
    # 四、阻尼牛顿法
    _add_heading(doc, 1, "四、阻尼牛顿求解算法")
    _add_heading(doc, 2, "（一）Newton方向与正权重步长")
    _add_body(
        doc,
        "在第k次迭代中，通过求解线性方程 Hₖpₖ=−gₖ 得到 Newton 方向，不显式计算 Hessian 的逆。由于 Hₖ 正定，"
        "梯度非零时有 gₖᵀpₖ<0。纯 Newton 步仍可能越出正权重域，因此先依据 pₖ 中的负分量计算最大可行步长，并乘以"
        "安全系数，使每次候选点都满足 xₖ+αpₖ>0。"
    )
    _add_equation(doc, equations, """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msub><mi>H</mi><mi>k</mi></msub><msub><mi>p</mi><mi>k</mi></msub><mo>=</mo><mo>−</mo><msub><mi>g</mi><mi>k</mi></msub><mo>,</mo><mspace width="0.5em"/><msub><mi>g</mi><mi>k</mi></msub><mo>=</mo><mo>∇</mo><mi>f</mi><mo stretchy="false">（</mo><msub><mi>x</mi><mi>k</mi></msub><mo stretchy="false">）</mo></mrow></math>""", 11)
    _add_heading(doc, 2, "（二）Armijo回溯与停止条件")
    _add_body(
        doc,
        "可行步长只保证变量为正，还需Armijo条件保证目标函数获得与方向导数相匹配的充分下降。若条件不成立，步长按固定比例回退。"
        "算法以无穷范数梯度、步长变化和最大迭代次数作为内部停止条件；归一化后重新计算RC误差和权重和误差，只有最大绝对RC"
        "偏差不超过10⁻⁶时才视为金融目标达标。"
    )
    _add_body(
        doc,
        "阻尼机制兼顾全局稳定性与局部速度：远离最优点时，回溯线搜索缩短过于激进的Newton步，使目标值沿下降方向稳定减少；进入"
        "最优点邻域后，单位步长通常能够直接被接受，并恢复Newton法的快速局部收敛。由于计算机中的梯度为小并不必然意味着归一化"
        "风险贡献已达标，本文将优化停止与经济约束验收分开记录。这一双重准则也使不同算法的比较不依赖各软件包对success状态的定义。"
    )
    _add_equation(doc, equations, """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mi>f</mi><mo stretchy="false">（</mo><msub><mi>x</mi><mi>k</mi></msub><mo>+</mo><mi>α</mi><msub><mi>p</mi><mi>k</mi></msub><mo stretchy="false">）</mo><mo>≤</mo><mi>f</mi><mo stretchy="false">（</mo><msub><mi>x</mi><mi>k</mi></msub><mo stretchy="false">）</mo><mo>+</mo><mi>c</mi><mi>α</mi><msubsup><mi>g</mi><mi>k</mi><mi>T</mi></msubsup><msub><mi>p</mi><mi>k</mi></msub></mrow></math>""", 12)
    _add_heading(doc, 2, "（三）算法流程与复杂度")
    for item in (
        "输入正定风险矩阵Σ、正风险预算b、容差和最大迭代次数，并以严格正向量初始化。",
        "计算目标函数、解析梯度和Hessian；梯度达到容差时进入解后检查。",
        "解线性方程得到Newton方向，并计算保持所有分量为正的最大步长。",
        "执行Armijo回溯，接受满足可行性和充分下降的候选点。",
        "将x归一化为资金权重w，输出权重、风险贡献、迭代诊断和独立RC验收结果。",
    ):
        _add_list_item(doc, item)
    _add_body(
        doc,
        "对n个资产，稠密Hessian的构造和线性方程求解通常为O(n³)，内存需求为O(n²)。本研究只有9类资产，解析二阶信息的"
        "成本较低；在超高维问题中，不显式存储Hessian的L-BFGS类方法可能更具扩展性[7]。"
    )

    # 五、数值实验
    _add_heading(doc, 1, "五、数值实验与算法比较")
    _add_heading(doc, 2, "（一）实验口径")
    _add_body(
        doc,
        "数值实验使用148个滚动月末EWMA半协方差矩阵，固定风险预算、初始化和1000次最大迭代上限。阻尼牛顿法与L-BFGS-B"
        "求解同一凸目标，SLSQP作为直接RC残差和显式资金约束下的通用算法基准。三者最终都按相同RC误差和权重约束重新验收；"
        "运行时间仅是当前软硬件环境中的辅助信息。"
    )
    worst_condition = float(stress["condition_number"].max())
    stress_row = stress.loc[(stress["condition_number"] == worst_condition) & (stress["method"] == "newton")].iloc[0]
    algorithm_rows: list[tuple[str, ...]] = []
    for method in ("newton", "lbfgsb", "slsqp"):
        row = algorithms.loc[algorithms["method"] == method].iloc[0]
        algorithm_rows.append((
            _method_label(method),
            _pct(row["success_rate"], 1),
            _num(row["median_iterations"], 1),
            _sci(row["median_rc_error"]),
            _sci(row["max_rc_error"]),
            _sci(row["max_weight_sum_error"]),
        ))
    algorithm_rows.append((
        "Newton：条件数10⁸",
        _pct(stress_row["success_rate"], 1),
        _num(stress_row["median_iterations"], 1),
        _sci(stress_row["median_rc_error"]),
        _sci(stress_row["max_rc_error"]),
        "—",
    ))
    algorithm_table = _add_table(
        doc,
        state,
        "求解算法与病态矩阵结果",
        ["方法/场景", "成功率", "中位迭代", "中位RC误差", "最大RC误差", "权重和误差"],
        algorithm_rows,
        [1.35, 0.75, 0.8, 1.08, 1.08, 1.12],
        font_size=8.8,
    )
    solver_figure = _add_figure(doc, state, figures_dir / "solver_summary.png", "三类求解算法的迭代效率与RC精度", width=5.9)
    _add_body_with_ref(
        doc,
        "算法汇总见表 ",
        algorithm_table,
        f"。阻尼牛顿法中位迭代 {newton['median_iterations']:.0f} 次，中位RC误差为 {_sci(newton['median_rc_error'])}，"
        "全部滚动窗口的最大RC误差低于10⁻⁶。L-BFGS-B无需完整Hessian但迭代次数较多；SLSQP的直接残差形式对停止设置更敏感。",
        str(state.table),
    )
    _add_body_with_ref(
        doc,
        "图形比较见图 ",
        solver_figure,
        "。在条件数10⁸的构造矩阵上，阻尼牛顿法的最大RC误差仍低于验收线，但成功率和迭代次数有所恶化，说明正则化、"
        "线搜索和解后诊断对病态输入具有实际必要性。",
        str(state.figure),
    )

    # 六、样本外结果
    _add_heading(doc, 1, "六、样本外实证结果")
    _add_heading(doc, 2, "（一）累计净值与核心绩效")
    _add_body(
        doc,
        "实证比较资金等权、逆下行波动率和ERC三类策略。三者采用相同调仓日、成本和可用数据，差异集中在权重生成规则。"
        "等权不使用风险估计，逆下行波动率忽略资产间相关性，ERC同时使用风险矩阵和目标风险预算。"
    )
    nav_figure = _add_figure(doc, state, figures_dir / "strategy_nav.png", "三类策略的累计净值", width=5.95)
    validation_rows = []
    for _, row in strategies.loc[strategies["period"] == "validation"].iterrows():
        validation_rows.append((
            _strategy_label(row["strategy"]),
            _pct(row["annual_return"]),
            _pct(row["annual_volatility"]),
            _num(row["sharpe"]),
            _pct(row["max_drawdown"]),
            _num(row["annual_turnover"]),
        ))
    validation_table = _add_table(
        doc,
        state,
        "验证期策略绩效",
        ["策略", "年化收益", "年化波动", "夏普比率", "最大回撤", "年换手"],
        validation_rows,
        [1.28, 0.9, 0.9, 0.8, 0.9, 1.4],
        font_size=9.2,
    )
    _add_body_with_ref(
        doc,
        "累计净值见图 ",
        nav_figure,
        "。等权组合在部分权益上涨阶段获得更高绝对收益，但净值波动和回撤也更大。风险平价的净值路径相对平滑，评价时应"
        "同时观察收益、波动、回撤和换手，而不能只比较终点净值。",
        str(state.figure),
    )
    _add_body_with_ref(
        doc,
        "验证期结果见表 ",
        validation_table,
        f"。ERC年化收益为 {_pct(validation_erc['annual_return'])}，低于等权的 {_pct(validation_equal['annual_return'])}；"
        f"但年化波动由 {_pct(validation_equal['annual_volatility'])} 降至 {_pct(validation_erc['annual_volatility'])}，"
        f"最大回撤由 {_pct(validation_equal['max_drawdown'])} 收窄至 {_pct(validation_erc['max_drawdown'])}，夏普比率提高至 "
        f"{_num(validation_erc['sharpe'])}。核心结论是风险调整后表现改善，而不是绝对收益最高。",
        str(state.table),
    )
    _add_body(
        doc,
        "上述差异应从风险暴露而非择时能力解释。滚动权重仅使用调仓日前已经观察到的收益，验证期参数亦未根据事后表现调整，因此"
        "回测避免了显式前视信息；但风险估计仍可能在市场结构突变时滞后。ERC降低波动和回撤，主要源于把资金从高边际风险资产转向"
        "低边际风险资产，并同时考虑相关性。相应代价是强势权益行情中可能落后于等权组合，且较高的低波动资产仓位也不等于不存在"
        "利率、流动性或尾部相关性风险。"
    )
    _add_heading(doc, 2, "（二）代表窗口的资金权重与风险贡献")
    representative_date = str(representative["representative_date"].iloc[0])
    weight_figure = _add_figure(doc, state, figures_dir / "weights_risk_contributions.png", f"{representative_date}的资金权重与风险贡献", width=5.95)
    _add_body_with_ref(
        doc,
        "代表窗口结果见图 ",
        weight_figure,
        "。各资产资金权重差异明显，但实际风险贡献都接近11.11%。债券和低波动资产需要较高资金权重，高波动权益与商品"
        "以较低仓位承担相同风险；相关性通过Σw进入边际风险贡献，因此ERC既不是资金等权，也不是简单逆波动率配置。",
        str(state.figure),
    )
    _add_heading(doc, 2, "（三）参数敏感性")
    sensitivity_figure = _add_figure(doc, state, figures_dir / "sensitivity_heatmap.png", "窗口长度和EWMA衰减系数的训练期敏感性", width=5.8)
    _add_body_with_ref(
        doc,
        "参数敏感性见图 ",
        sensitivity_figure,
        f"。按训练期夏普最高、再选择最大回撤较轻的规则，网格选中窗口 {selected['window']}、衰减系数 "
        f"{selected['decay']:.2f}，训练期夏普为 {_num(selected['train_sharpe'])}，验证期夏普为 "
        f"{_num(selected['validation_sharpe'])}。主模型仍使用项目既定的252日和0.97参数，以保持预设口径；敏感性结果只用于"
        "说明权重反应速度与估计噪声之间的权衡，不在验证期重新调参。",
        str(state.figure),
    )

    # 七、一般风险预算
    _add_heading(doc, 1, "七、一般风险预算扩展")
    _add_heading(doc, 2, "（一）模型推广")
    _add_body(
        doc,
        "当bᵢ不再相等时，式（7）—（10）仍然成立，阻尼牛顿法无需改变算法结构。预算提高意味着资产承担更高风险份额，"
        "但最终资金权重还取决于自身波动和与其他资产的相关性。在对角协方差下可得到闭式关系wᵢ∝√(bᵢ/Σᵢᵢ)，"
        "该关系也作为一般风险预算实现的单元测试。"
    )
    _add_equation(doc, equations, """<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msub><mi>w</mi><mi>i</mi></msub><mo>∝</mo><msqrt><mfrac><msub><mi>b</mi><mi>i</mi></msub><msub><mi>Σ</mi><mrow><mi>i</mi><mi>i</mi></mrow></msub></mfrac></msqrt></mrow></math>""", 13)
    _add_heading(doc, 2, "（二）非等额预算示例")
    rb_rows = []
    for _, row in risk_budget.iterrows():
        rb_rows.append((
            row["asset"],
            _num(row["raw_budget_multiplier"], 1),
            _pct(row["target_risk_budget"]),
            _pct(row["actual_risk_contribution"]),
            _pct(row["weight"]),
            _sci(row["absolute_rc_error"]),
        ))
    rb_table = _add_table(
        doc,
        state,
        f"{rb_summary['representative_date']}一般风险预算代表解",
        ["资产", "预算倍率", "目标预算", "实际RC", "资金权重", "RC误差"],
        rb_rows,
        [1.32, 0.82, 0.9, 0.9, 0.9, 1.34],
        font_size=8.8,
    )
    rb_figure = _add_figure(doc, state, figures_dir / "risk_budget_extension.png", "一般风险预算的目标贡献、实际贡献与资金权重", width=5.9)
    _add_body_with_ref(
        doc,
        "代表实验将沪深300ETF和中证1000ETF的预算倍率设为2，其余资产为1。归一化后，两类权益资产的目标风险预算均为"
        "18.18%，其他资产为9.09%。详细结果见表 ",
        rb_table,
        f"，最大预算跟踪误差为 {_sci(rb_summary['rc_max_error'])}。",
        str(state.table),
    )
    _add_body_with_ref(
        doc,
        "目标预算、实际风险贡献和资金权重的关系见图 ",
        rb_figure,
        "。实际RC几乎与目标重合，说明同一凸模型可以准确实现非等额预算；资金权重则不会简单按预算倍率同比变化。"
        "该实验只证明模型推广和数值求解能力，不意味着提高权益风险预算能够改善未来收益。",
        str(state.figure),
    )

    # 八、结论与局限
    _add_heading(doc, 1, "八、结论与局限")
    _add_heading(doc, 2, "（一）主要结论")
    _add_body(
        doc,
        "第一，风险预算条件可以转化为正权重域上的严格凸问题，解析梯度、Hessian和KKT条件共同保证唯一解及风险贡献等价关系。"
        "第二，带正权重步长和Armijo回溯的阻尼牛顿法能够利用问题的二阶结构，以较少迭代达到高精度；求解器内部状态仍需由RC"
        "误差和权重约束独立复核。第三，样本外结果显示ERC主要降低波动和最大回撤，提高风险调整后收益，但不保证最高绝对收益。"
        "第四，一般风险预算可以表达非等额风险偏好，且无需改变凸目标和核心求解流程。"
    )
    _add_heading(doc, 2, "（二）研究局限")
    _add_body(
        doc,
        "本研究仍有四点局限：ETF成立前使用指数代理，未完整刻画真实可交易性和跟踪误差；风险矩阵依赖有限窗口和EWMA参数，"
        "尚未系统比较收缩估计或因子模型；回测采用固定5bp成本，未覆盖冲击成本和容量约束；样本外结果来自特定资产池和市场时期，"
        "不能外推为未来收益保证。后续可在保持凸结构的前提下研究协方差收缩、权重上下限、换手惩罚和滚动交叉验证。"
    )

    # 参考文献
    _add_heading(doc, 1, "参考文献")
    references = [
        "[1] MARKOWITZ H. Portfolio selection[J]. The Journal of Finance, 1952, 7(1): 77-91.",
        "[2] TASCHE D. Capital allocation to business units and sub-portfolios: The Euler principle[EB/OL]. arXiv:0708.2542, 2008.",
        "[3] MAILLARD S, RONCALLI T, TEÏLETCHE J. On the properties of equally-weighted risk contributions portfolios[J]. The Journal of Portfolio Management, 2010, 36(4): 60-70.",
        "[4] BRUDER B, RONCALLI T. Managing risk exposures using the risk budgeting approach[R]. MPRA Paper No. 37749, 2012.",
        "[5] CETINGOZ A R, FERMANIAN J D, GUÉANT O. Risk budgeting portfolios: Existence and computation[EB/OL]. arXiv:2211.07212, 2023.",
        "[6] SPINU F. An algorithm for computing risk parity weights[R]. SSRN Working Paper 2297383, 2013.",
        "[7] NOCEDAL J, WRIGHT S J. Numerical Optimization[M]. 2nd ed. New York: Springer, 2006.",
        "[8] RONCALLI T. Introduction to Risk Parity and Budgeting[M]. Boca Raton: CRC Press, 2013.",
        "[9] J.P. MORGAN, REUTERS. RiskMetrics Technical Document[R]. 4th ed. New York: Morgan Guaranty Trust Company, 1996.",
    ]
    for reference in references:
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.paragraph_format.first_line_indent = Pt(-21)
        paragraph.paragraph_format.left_indent = Pt(21)
        paragraph.paragraph_format.line_spacing = 1.2
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(reference)
        _set_run_font(run, size=10.5)


def _run_word_export(course_dir: Path, docx_path: Path, pdf_path: Path) -> dict[str, Any]:
    script = Path(__file__).with_name("export_word_pdf.ps1")
    result_path = course_dir / "tmp" / "word_report" / "word_export.json"
    result_path.parent.mkdir(parents=True, exist_ok=True)
    command = [
        "powershell.exe",
        "-NoProfile",
        "-ExecutionPolicy",
        "Bypass",
        "-File",
        str(script),
        "-DocxPath",
        str(docx_path),
        "-PdfPath",
        str(pdf_path),
        "-ResultJson",
        str(result_path),
    ]
    process = subprocess.run(
        command,
        cwd=course_dir,
        capture_output=True,
        text=True,
        errors="replace",
        timeout=180,
    )
    if process.returncode != 0:
        raise RuntimeError(
            "Microsoft Word 导出失败。\n"
            f"STDOUT:\n{process.stdout}\nSTDERR:\n{process.stderr}"
        )
    if not result_path.exists() or not pdf_path.exists() or pdf_path.stat().st_size < 10_000:
        raise RuntimeError("Microsoft Word 未生成有效 PDF")
    return json.loads(result_path.read_text(encoding="utf-8-sig"))


def _scrub_final_docx(docx_path: Path, config: dict[str, Any]) -> None:
    temporary = docx_path.with_suffix(".scrubbed.docx")
    with ZipFile(docx_path, "r") as source, ZipFile(temporary, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            name = item.filename
            if name == "docProps/custom.xml":
                continue
            data = source.read(name)
            if name == "docProps/core.xml":
                root = etree.fromstring(data)
                values = {
                    f"{{{DC_NS}}}title": str(config["title"]),
                    f"{{{DC_NS}}}subject": "最优化理论与算法课程报告",
                    f"{{{DC_NS}}}creator": "",
                    f"{{{CP_NS}}}lastModifiedBy": "",
                    f"{{{CP_NS}}}keywords": "风险平价；凸优化；阻尼牛顿法；风险预算；资产配置",
                    f"{{{DC_NS}}}description": "由项目数据、模型推导与实验结果可复现生成",
                }
                for tag, value in values.items():
                    element = root.find(tag)
                    if element is not None:
                        element.text = value
                data = _xml_bytes(root)
            elif name.startswith("word/") and name.endswith(".xml"):
                try:
                    root = etree.fromstring(data)
                except etree.XMLSyntaxError:
                    pass
                else:
                    for element in root.iter():
                        for attr in list(element.attrib):
                            if attr.startswith(f"{{{WORD_NS}}}rsid"):
                                del element.attrib[attr]
                    data = _xml_bytes(root)
            target.writestr(item, data)
    temporary.replace(docx_path)


def _docx_xml_text(docx_path: Path) -> tuple[str, str]:
    with ZipFile(docx_path, "r") as archive:
        xml_parts = [
            name
            for name in archive.namelist()
            if name.endswith(".xml") and (name.startswith("word/") or name.startswith("docProps/"))
        ]
        raw_xml = "\n".join(archive.read(name).decode("utf-8", errors="ignore") for name in xml_parts)
        root = etree.fromstring(archive.read("word/document.xml"))
        text = "".join(root.itertext())
    return text, raw_xml


def _verify_word_outputs(
    docx_path: Path,
    pdf_path: Path,
    word_export: dict[str, Any],
    config: dict[str, Any],
    summary: dict[str, Any],
) -> dict[str, Any]:
    text, raw_xml = _docx_xml_text(docx_path)
    pdf_reader = PdfReader(str(pdf_path))
    pdf_text = "\n".join((page.extract_text() or "") for page in pdf_reader.pages)
    required_terms = [
        "凸等价重构",
        "阻尼牛顿法",
        "Armijo回溯",
        "一般风险预算",
        "风险贡献",
        "参考文献",
    ]
    forbidden_terms = [
        "程思宇",
        "2025212590",
        "地缘政治冲突冲击下全球多资产组合的尾部风险约束优化研究",
        "优化器演进",
        "目标放大",
        "受控复现",
        "v0.02",
        "v0.03",
        "v0.04",
        "v0.05",
        "v0.06",
        "v0.15",
        "v0.16_2",
        "v0.16_3",
    ]
    compact_docx_text = re.sub(r"\s+", "", text)
    compact_pdf_text = re.sub(r"\s+", "", pdf_text)
    missing_docx = [term for term in required_terms if term not in compact_docx_text]
    missing_pdf = [term for term in required_terms if term not in compact_pdf_text]
    forbidden_found = [term for term in forbidden_terms if term in text or term in raw_xml]
    field_counts = {
        name: len(re.findall(rf"\b{name}\b", raw_xml))
        for name in ("TOC", "PAGE", "SEQ", "REF")
    }
    equation_count = raw_xml.count("<m:oMath")
    a4_pages = all(
        abs(float(page.mediabox.width) - 595.3) < 2.0
        and abs(float(page.mediabox.height) - 841.9) < 2.0
        for page in pdf_reader.pages
    )
    pdf_base_fonts: set[str] = set()
    for page in pdf_reader.pages:
        resources = page.get("/Resources", {})
        fonts = resources.get("/Font", {}) if resources else {}
        for reference in fonts.values():
            font = reference.get_object()
            base_font = font.get("/BaseFont")
            if base_font:
                pdf_base_fonts.add(str(base_font))
    unexpected_title_fonts = sorted(
        font
        for font in pdf_base_fonts
        if any(marker in font for marker in ("MicrosoftYaHei", "MS-Gothic", "Calibri-Bold"))
    )
    key_numbers = [
        _pct(summary["validation_erc"]["annual_return"]),
        _pct(summary["validation_erc"]["annual_volatility"]),
        _num(summary["validation_erc"]["sharpe"]),
        _sci(summary["newton_summary"]["median_rc_error"]),
        "18.18%",
    ]
    missing_numbers = [value for value in key_numbers if value not in text]
    page_count = len(pdf_reader.pages)
    word_pages = int(word_export["page_count"])
    document = Document(docx_path)
    def effective_style_font(style: Any, attribute: str) -> str | None:
        current = style
        while current is not None:
            r_pr = current.element.rPr
            if r_pr is not None and r_pr.rFonts is not None:
                value = r_pr.rFonts.get(qn(attribute))
                if value:
                    return value
            current = current.base_style
        return None

    title_style_fonts = {
        name: effective_style_font(document.styles[name], "w:eastAsia")
        for name in ("Heading 1", "Heading 2", "Heading 3", "Caption")
    }
    title_fonts_are_simsun = all(value == "宋体" for value in title_style_fonts.values())
    checks = {
        "docx_exists": docx_path.exists() and docx_path.stat().st_size > 50_000,
        "pdf_exists": pdf_path.exists() and pdf_path.stat().st_size > 50_000,
        "export_engine_is_word": word_export.get("export_engine") == "Microsoft Word",
        "word_pdf_page_count_matches": page_count == word_pages,
        "page_count_within_reasonable_range": 10 <= page_count <= 17,
        "all_pages_a4": a4_pages,
        "selectable_pdf_text": len(pdf_text) > 8_000,
        "required_terms_present": not missing_docx and not missing_pdf,
        "forbidden_reference_content_absent": not forbidden_found,
        "key_numbers_present": not missing_numbers,
        "toc_present": field_counts["TOC"] >= 1,
        "page_fields_present": field_counts["PAGE"] >= 2,
        "caption_fields_present": field_counts["SEQ"] >= 9,
        "cross_reference_fields_present": field_counts["REF"] >= 5,
        "native_equations_present": equation_count >= 12,
        "title_fonts_are_simsun": title_fonts_are_simsun,
        "word_pdf_has_no_fallback_title_fonts": not unexpected_title_fonts,
    }
    return {
        "docx": str(docx_path),
        "pdf": str(pdf_path),
        "source_docx": str(docx_path),
        "export_engine": word_export["export_engine"],
        "word_version": word_export["word_version"],
        "word_statistics": {
            "page_count": word_pages,
            "word_count": int(word_export["word_count"]),
            "character_count": int(word_export["character_count"]),
        },
        "file_size_bytes": pdf_path.stat().st_size,
        "docx_size_bytes": docx_path.stat().st_size,
        "docx_sha256": _sha256(docx_path),
        "pdf_sha256": _sha256(pdf_path),
        "page_count": page_count,
        "page_size_points": [
            float(pdf_reader.pages[0].mediabox.width),
            float(pdf_reader.pages[0].mediabox.height),
        ],
        "selectable_text_characters": len(pdf_text),
        "required_terms": required_terms,
        "missing_required_terms_docx": missing_docx,
        "missing_required_terms_pdf": missing_pdf,
        "forbidden_terms_found": forbidden_found,
        "missing_key_numbers": missing_numbers,
        "field_counts": field_counts,
        "native_equation_count": equation_count,
        "title_style_fonts": title_style_fonts,
        "pdf_base_fonts": sorted(pdf_base_fonts),
        "unexpected_title_fonts": unexpected_title_fonts,
        "checks": checks,
        "personal_info_pending": config.get("student_name") == "请填写" or config.get("student_id") == "请填写",
        "status": "passed" if all(checks.values()) else "needs_revision",
    }


def build_word_and_pdf(course_dir: Path, config: dict[str, Any], summary: dict[str, Any]) -> dict[str, Any]:
    template_path = course_dir / "templates" / "course_report_template.docx"
    if not template_path.exists():
        raise FileNotFoundError(
            f"缺少清洁 Word 模板：{template_path}。模板必须先由参考 Word 脱敏生成，构建过程不回退到下载目录。"
        )

    docx_dir = course_dir / "output" / "docx"
    pdf_dir = course_dir / "output" / "pdf"
    docx_dir.mkdir(parents=True, exist_ok=True)
    pdf_dir.mkdir(parents=True, exist_ok=True)
    docx_path = docx_dir / "最优化理论与算法期末大作业_风险平价.docx"
    pdf_path = pdf_dir / "最优化理论与算法期末大作业_风险平价.pdf"

    doc = Document(template_path)
    _clear_document_body(doc)
    _configure_styles(doc)
    _set_document_properties(doc, config)
    _set_update_fields(doc)

    cover_section = doc.sections[0]
    _configure_section(cover_section)
    cover_section.footer.is_linked_to_previous = False
    _clear_story(cover_section.footer)
    _add_cover(doc, config)

    front_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(front_section)
    _add_page_number(front_section, "lowerRoman", 1)
    _add_abstract(doc, config, summary)
    doc.add_page_break()
    _add_front_title(doc, "目　录")
    _add_toc(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(body_section)
    _add_page_number(body_section, "decimal", 1)
    _add_report_content(doc, course_dir, config, summary)

    doc.save(docx_path)
    word_export = _run_word_export(course_dir, docx_path, pdf_path)
    _scrub_final_docx(docx_path, config)
    verification = _verify_word_outputs(docx_path, pdf_path, word_export, config, summary)
    verification_path = docx_dir / "word_report_verification.json"
    verification_path.write_text(json.dumps(verification, ensure_ascii=False, indent=2), encoding="utf-8")
    if verification["status"] != "passed":
        raise RuntimeError(f"Word/PDF 结构验收未通过：{verification}")
    return verification
