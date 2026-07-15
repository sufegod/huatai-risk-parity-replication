from __future__ import annotations

from dataclasses import replace
import json
from pathlib import Path
import re
import sys
import unittest
from zipfile import ZipFile

import numpy as np
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
from lxml import html as lxml_html
from pypdf import PdfReader


COURSE_DIR = Path(__file__).resolve().parents[1]
REPO_DIR = COURSE_DIR.parent
if str(COURSE_DIR) not in sys.path:
    sys.path.insert(0, str(COURSE_DIR))

from src.backtest import BacktestConfig, run_backtest
from src.optimizer_evolution import (
    OPTIMIZER_VARIANTS,
    evaluate_optimizer_evolution,
    historical_convex_gradient,
    historical_convex_objective,
    relative_erc_objective,
    scaled_erc_objective,
    solve_optimizer_variant,
)
from src.risk_parity import (
    convex_gradient,
    convex_hessian,
    convex_objective,
    estimate_covariance,
    original_erc_objective,
    risk_contributions,
    solve_erc,
)


class RiskParityModelTests(unittest.TestCase):
    def setUp(self):
        self.covariance = np.array(
            [
                [0.040, 0.006, 0.002],
                [0.006, 0.090, 0.004],
                [0.002, 0.004, 0.160],
            ],
            dtype=float,
        )
        self.budget = np.full(3, 1.0 / 3.0)

    def test_analytic_gradient_matches_finite_difference(self):
        x = np.array([0.8, 1.1, 1.4])
        step = 1e-6
        numerical = np.empty_like(x)
        for i in range(len(x)):
            upper = x.copy()
            lower = x.copy()
            upper[i] += step
            lower[i] -= step
            numerical[i] = (
                convex_objective(upper, self.covariance, self.budget)
                - convex_objective(lower, self.covariance, self.budget)
            ) / (2 * step)
        np.testing.assert_allclose(convex_gradient(x, self.covariance, self.budget), numerical, rtol=1e-6, atol=1e-8)

    def test_analytic_hessian_matches_gradient_difference_and_is_spd(self):
        x = np.array([0.8, 1.1, 1.4])
        step = 1e-6
        numerical = np.empty((3, 3))
        for i in range(len(x)):
            upper = x.copy()
            lower = x.copy()
            upper[i] += step
            lower[i] -= step
            numerical[:, i] = (
                convex_gradient(upper, self.covariance, self.budget)
                - convex_gradient(lower, self.covariance, self.budget)
            ) / (2 * step)
        analytic = convex_hessian(x, self.covariance, self.budget)
        np.testing.assert_allclose(analytic, numerical, rtol=1e-6, atol=1e-8)
        self.assertGreater(np.linalg.eigvalsh(analytic).min(), 0.0)

    def test_diagonal_covariance_matches_inverse_volatility_closed_form(self):
        diagonal = np.diag(np.square([0.10, 0.20, 0.40, 0.80]))
        expected = 1.0 / np.sqrt(np.diag(diagonal))
        expected /= expected.sum()
        for method in ("newton", "lbfgsb", "slsqp"):
            result = solve_erc(diagonal, method=method, tol=1e-10, max_iter=1000)
            np.testing.assert_allclose(result.weights, expected, atol=2e-6)
            self.assertLess(result.rc_max_error, 1e-6)

    def test_general_risk_budget_matches_diagonal_closed_form(self):
        diagonal = np.diag(np.square([0.10, 0.20, 0.40, 0.80]))
        budget = np.array([0.4, 0.3, 0.2, 0.1])
        expected = np.sqrt(budget / np.diag(diagonal))
        expected /= expected.sum()
        result = solve_erc(diagonal, method="newton", risk_budget=budget, tol=1e-10)
        np.testing.assert_allclose(result.weights, expected, atol=2e-8)
        np.testing.assert_allclose(risk_contributions(result.weights, diagonal), budget, atol=1e-8)

    def test_solvers_agree_and_obey_constraints(self):
        results = {method: solve_erc(self.covariance, method=method) for method in ("newton", "lbfgsb", "slsqp")}
        for result in results.values():
            self.assertAlmostEqual(float(result.weights.sum()), 1.0, places=12)
            self.assertGreater(result.weights.min(), 0.0)
            self.assertLess(result.rc_max_error, 1e-6)
        np.testing.assert_allclose(results["newton"].weights, results["lbfgsb"].weights, atol=2e-6)
        np.testing.assert_allclose(results["newton"].weights, results["slsqp"].weights, atol=2e-6)

    def test_newton_is_deterministic(self):
        first = solve_erc(self.covariance, method="newton")
        second = solve_erc(self.covariance, method="newton")
        np.testing.assert_array_equal(first.weights, second.weights)
        self.assertEqual(first.iterations, second.iterations)

    def test_ewma_semi_covariance_is_positive_definite(self):
        rng = np.random.default_rng(42)
        returns = rng.normal(0.0, 0.01, size=(252, 9))
        covariance = estimate_covariance(returns, method="ewma_semi", decay=0.97, ridge=1e-8)
        self.assertGreater(np.linalg.eigvalsh(covariance).min(), 0.0)


class OptimizerEvolutionTests(unittest.TestCase):
    def setUp(self):
        self.covariance = np.array(
            [
                [0.040, 0.006, 0.002],
                [0.006, 0.090, 0.004],
                [0.002, 0.004, 0.160],
            ],
            dtype=float,
        )
        self.weights = np.array([0.25, 0.35, 0.40])
        self.budget = np.full(3, 1.0 / 3.0)

    def test_absolute_objective_changes_quadratically_with_covariance_scale(self):
        base = original_erc_objective(self.weights, self.covariance, self.budget)
        scaled = original_erc_objective(self.weights, self.covariance * 7.0, self.budget)
        self.assertAlmostEqual(scaled, base * 49.0, places=12)

    def test_v003_objective_is_exactly_original_times_one_billion(self):
        base = original_erc_objective(self.weights, self.covariance, self.budget)
        self.assertAlmostEqual(scaled_erc_objective(self.weights, self.covariance), base * 1e9, places=8)

    def test_v004_relative_objective_is_scale_invariant(self):
        base = relative_erc_objective(self.weights, self.covariance)
        scaled = relative_erc_objective(self.weights, self.covariance * 13.0)
        self.assertAlmostEqual(base, scaled, places=12)

    def test_v005_analytic_gradient_matches_finite_difference(self):
        x = np.array([0.8, 1.1, 1.4])
        step = 1e-6
        numerical = np.empty_like(x)
        for i in range(len(x)):
            upper = x.copy()
            lower = x.copy()
            upper[i] += step
            lower[i] -= step
            numerical[i] = (
                historical_convex_objective(upper, self.covariance)
                - historical_convex_objective(lower, self.covariance)
            ) / (2 * step)
        np.testing.assert_allclose(
            historical_convex_gradient(x, self.covariance), numerical, rtol=1e-6, atol=1e-8
        )

    def test_all_controlled_variants_return_feasible_weights(self):
        for variant in OPTIMIZER_VARIANTS:
            with self.subTest(variant=variant):
                result = solve_optimizer_variant(self.covariance, variant)
                self.assertAlmostEqual(float(result.weights.sum()), 1.0, places=10)
                self.assertGreaterEqual(result.min_weight, -1e-12)

    def test_full_controlled_experiment_contains_740_diagnostics(self):
        returns = pd.read_csv(COURSE_DIR / "data" / "etf_returns.csv", index_col=0, parse_dates=True)
        details, summary = evaluate_optimizer_evolution(returns)
        self.assertEqual(len(details), 740)
        self.assertEqual(len(summary), 5)
        self.assertTrue((details.groupby("variant").size() == 148).all())
        self.assertGreaterEqual(float(details["min_weight"].min()), -1e-12)
        accepted = details.loc[details["rc_pass"]]
        self.assertLessEqual(float(accepted["weight_sum_error"].max()), 1e-10)
        current = details.loc[details["variant"] == "course_newton"]
        self.assertTrue(current["rc_pass"].all())
        self.assertLessEqual(float(current["weight_sum_error"].max()), 1e-10)


class BacktestTests(unittest.TestCase):
    @staticmethod
    def synthetic_returns() -> pd.DataFrame:
        rng = np.random.default_rng(7)
        dates = pd.bdate_range("2018-01-01", periods=700)
        values = rng.normal(0.0002, [0.008, 0.012, 0.004], size=(700, 3))
        return pd.DataFrame(values, index=dates, columns=["A", "B", "C"])

    def config(self, **changes) -> BacktestConfig:
        base = BacktestConfig(
            strategy="erc",
            window=126,
            decay=0.97,
            fee_rate=0.0005,
            train_start="2018-01-01",
            train_end="2019-12-31",
            validation_start="2020-01-01",
            validation_end="2021-12-31",
        )
        return replace(base, **changes)

    def test_future_changes_do_not_affect_first_target(self):
        returns = self.synthetic_returns()
        original = run_backtest(returns, self.config())
        changed = returns.copy()
        first_execution = original.target_weights.index[0]
        changed.loc[changed.index > first_execution] *= 5.0
        rerun = run_backtest(changed, self.config())
        np.testing.assert_allclose(
            original.target_weights.iloc[0].to_numpy(),
            rerun.target_weights.iloc[0].to_numpy(),
            atol=1e-12,
        )

    def test_transaction_cost_reduces_nav(self):
        returns = self.synthetic_returns() * 0.0
        without_cost = run_backtest(returns, self.config(strategy="equal_weight", fee_rate=0.0))
        with_cost = run_backtest(returns, self.config(strategy="equal_weight", fee_rate=0.0005))
        self.assertLess(with_cost.nav.iloc[-1], without_cost.nav.iloc[-1])
        self.assertAlmostEqual(without_cost.nav.iloc[-1], 1.0, places=12)

    def test_rebalance_executes_after_observation(self):
        result = run_backtest(self.synthetic_returns(), self.config())
        observations = pd.to_datetime(result.solver_diagnostics["observation_date"])
        self.assertTrue((result.solver_diagnostics.index > observations).all())


class SourceDataTests(unittest.TestCase):
    def test_source_data_shape_and_quality(self):
        source = REPO_DIR / "数据" / "原始数据" / "ETF风险平价回测数据.xlsx"
        frame = pd.read_excel(source, sheet_name=0, index_col=0, parse_dates=True)
        self.assertEqual(frame.shape, (3216, 9))
        self.assertEqual(int(frame.index.duplicated().sum()), 0)
        self.assertEqual(int(frame.isna().sum().sum()), 3)
        self.assertEqual(str(frame.index.min().date()), "2013-01-04")
        self.assertEqual(str(frame.index.max().date()), "2026-04-03")

    def test_newton_converges_quickly_on_representative_real_window(self):
        source = REPO_DIR / "数据" / "原始数据" / "ETF风险平价回测数据.xlsx"
        frame = pd.read_excel(source, sheet_name=0, index_col=0, parse_dates=True).fillna(0.0) / 100.0
        covariance = estimate_covariance(frame.iloc[-252:], method="ewma_semi", decay=0.97, ridge=1e-8)
        result = solve_erc(covariance, method="newton", tol=1e-10)
        self.assertLessEqual(result.iterations, 10)
        self.assertLess(result.rc_max_error, 1e-6)
        np.testing.assert_allclose(risk_contributions(result.weights, covariance), np.full(9, 1 / 9), atol=1e-6)


class WordReportTests(unittest.TestCase):
    DOCX_PATH = COURSE_DIR / "output" / "docx" / "最优化理论与算法期末大作业_风险平价.docx"
    PDF_PATH = COURSE_DIR / "output" / "pdf" / "最优化理论与算法期末大作业_风险平价.pdf"
    TEMPLATE_PATH = COURSE_DIR / "templates" / "course_report_template.docx"

    @staticmethod
    def _docx_payload(path: Path) -> tuple[str, str, list[str]]:
        with ZipFile(path) as archive:
            names = archive.namelist()
            xml = "\n".join(
                archive.read(name).decode("utf-8", errors="ignore")
                for name in names
                if name.endswith(".xml")
            )
            document = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        return document, xml, names

    def test_sanitized_template_preserves_layout_without_reference_content(self):
        self.assertTrue(self.TEMPLATE_PATH.exists())
        _, xml, names = self._docx_payload(self.TEMPLATE_PATH)
        for forbidden in ("程思宇", "2025212590", "CVaR", "地缘政治冲突冲击"):
            self.assertNotIn(forbidden, xml)
        self.assertNotIn("2013-12-23", xml)
        self.assertNotIn("<dcterms:created", xml)
        self.assertFalse(any(name.startswith("word/media/") for name in names))

        document = Document(self.TEMPLATE_PATH)
        section = document.sections[0]
        self.assertAlmostEqual(section.page_width / Cm(1), 21.0, places=1)
        self.assertAlmostEqual(section.page_height / Cm(1), 29.7, places=1)
        self.assertAlmostEqual(section.left_margin / Cm(1), 2.69, places=1)
        self.assertAlmostEqual(section.right_margin / Cm(1), 2.41, places=1)
        self.assertAlmostEqual(section.top_margin / Cm(1), 2.41, places=1)
        self.assertAlmostEqual(section.bottom_margin / Cm(1), 2.21, places=1)
        self.assertAlmostEqual(document.styles["Normal"].font.size.pt, 11.0, places=1)
        self.assertEqual(str(document.styles["Heading 1"].font.color.rgb), "365F91")
        self.assertEqual(str(document.styles["Heading 2"].font.color.rgb), "4F81BD")

    def test_word_report_has_required_structure_fields_and_native_equations(self):
        self.assertTrue(self.DOCX_PATH.exists())
        document_xml, xml, _ = self._docx_payload(self.DOCX_PATH)
        document = Document(self.DOCX_PATH)
        visible_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        visible_text += "\n" + "\n".join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        normalized_visible_text = re.sub(r"\s+", "", visible_text.replace("\u3000", ""))
        required_terms = (
            "摘要",
            "关键词",
            "凸等价重构",
            "阻尼牛顿法",
            "Armijo回溯",
            "一般风险预算",
            "参考文献",
        )
        for term in required_terms:
            self.assertIn(term, normalized_visible_text)
        self.assertIn("姓名：程哲", visible_text)
        self.assertIn("学号：2025212591", visible_text)
        for forbidden in (
            "程思宇",
            "2025212590",
            "地缘政治冲突冲击下全球多资产组合的尾部风险约束优化研究",
            "优化器演进",
            "目标放大",
            "受控复现",
            "v0.02",
            "v0.03",
            "v0.04",
            "v0.05",
            "v0.06",
            "v0.15",
            "v0.16_2",
            "v0.16_3",
        ):
            self.assertNotIn(forbidden, normalized_visible_text)
        self.assertNotIn("2013-12-23", xml)
        with ZipFile(self.DOCX_PATH) as archive:
            core_properties = archive.read("docProps/core.xml").decode("utf-8")
        self.assertIn("基于 EWMA 半协方差的风险平价资产配置优化", core_properties)
        self.assertNotIn("»ùÓÚ", core_properties)
        self.assertIn("TOC \\o", xml)
        self.assertGreaterEqual(len(re.findall(r"\bSEQ\b", xml)), 9)
        self.assertGreaterEqual(len(re.findall(r"\bREF\b", xml)), 5)
        self.assertGreaterEqual(xml.count("<m:oMath"), 12)
        self.assertIn("w:fldCharType=\"begin\"", document_xml)

        self.assertGreaterEqual(len(document.sections), 3)

        def effective_style_font(style, attribute):
            current = style
            while current is not None:
                r_pr = current.element.rPr
                if r_pr is not None and r_pr.rFonts is not None:
                    value = r_pr.rFonts.get(qn(attribute))
                    if value:
                        return value
                current = current.base_style
            return None

        for style_name in ("Heading 1", "Heading 2", "Heading 3", "Caption"):
            style = document.styles[style_name]
            self.assertEqual(effective_style_font(style, "w:eastAsia"), "宋体")
            self.assertEqual(effective_style_font(style, "w:ascii"), "Times New Roman")
        direct_title_texts = {
            "最优化理论与算法课程报告",
            "基于 EWMA 半协方差的风险平价资产配置优化——凸重构、阻尼牛顿法与实证分析",
            "摘　要",
            "目　录",
        }
        direct_titles = [
            p for p in document.paragraphs
            if p.text.strip().replace("\n", "") in direct_title_texts
        ]
        self.assertEqual(len(direct_titles), len(direct_title_texts))
        for paragraph in direct_titles:
            for run in paragraph.runs:
                if not run.text.strip():
                    continue
                r_fonts = run._element.get_or_add_rPr().rFonts
                east_asia = r_fonts.get(qn("w:eastAsia")) if r_fonts is not None else None
                ascii_font = r_fonts.get(qn("w:ascii")) if r_fonts is not None else None
                self.assertEqual(east_asia or effective_style_font(paragraph.style, "w:eastAsia"), "宋体")
                self.assertEqual(ascii_font or effective_style_font(paragraph.style, "w:ascii"), "Times New Roman")
        for section in document.sections:
            self.assertAlmostEqual(section.page_width / Cm(1), 21.0, places=1)
            self.assertAlmostEqual(section.page_height / Cm(1), 29.7, places=1)

    def test_word_pdf_engine_page_geometry_and_numeric_consistency(self):
        self.assertTrue(self.PDF_PATH.exists())
        verification = json.loads(
            (COURSE_DIR / "output" / "docx" / "word_report_verification.json").read_text(encoding="utf-8")
        )
        self.assertEqual(verification["status"], "passed")
        self.assertEqual(verification["export_engine"], "Microsoft Word")
        self.assertTrue(all(verification["checks"].values()))
        self.assertEqual(verification["unexpected_title_fonts"], [])

        reader = PdfReader(str(self.PDF_PATH))
        self.assertEqual(len(reader.pages), verification["word_statistics"]["page_count"])
        self.assertGreaterEqual(len(reader.pages), 10)
        self.assertLessEqual(len(reader.pages), 17)
        self.assertTrue(
            all(
                abs(float(page.mediabox.width) - 595.3) < 2.0
                and abs(float(page.mediabox.height) - 841.9) < 2.0
                for page in reader.pages
            )
        )
        pdf_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        _, docx_xml, _ = self._docx_payload(self.DOCX_PATH)
        html = (COURSE_DIR / "output" / "html" / "report.html").read_text(encoding="utf-8")
        html_visible_text = lxml_html.fromstring(html).text_content()
        summary = json.loads((COURSE_DIR / "output" / "tables" / "analysis_summary.json").read_text(encoding="utf-8"))

        annual_return = f"{summary['validation_erc']['annual_return']:.2%}"
        annual_volatility = f"{summary['validation_erc']['annual_volatility']:.2%}"
        sharpe = f"{summary['validation_erc']['sharpe']:.2f}"
        for value in (annual_return, annual_volatility, sharpe):
            self.assertIn(value, docx_xml)
            self.assertIn(value, pdf_text)
            self.assertIn(value, html)
        rc_error = f"{summary['newton_summary']['median_rc_error']:.2e}"
        self.assertIn(rc_error, docx_xml)
        self.assertIn(rc_error, pdf_text)
        self.assertIn(rc_error, html_visible_text)
        for forbidden in ("优化器演进", "目标放大", "受控复现", "v0.02", "v0.16_3"):
            self.assertNotIn(forbidden, docx_xml)
            self.assertNotIn(forbidden, pdf_text)
            self.assertNotIn(forbidden, html)

        risk_budget = pd.read_csv(COURSE_DIR / "output" / "tables" / "risk_budget_extension.csv")
        doubled = risk_budget.loc[risk_budget["raw_budget_multiplier"] == 2.0, "target_risk_budget"]
        self.assertEqual(len(doubled), 2)
        self.assertTrue(np.allclose(doubled.to_numpy(), 2.0 / 11.0))
        self.assertIn("18.18%", docx_xml)
        self.assertIn("18.18%", pdf_text)
        self.assertIn("18.2%", html)


if __name__ == "__main__":
    unittest.main()

